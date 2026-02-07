from fastapi import FastAPI, Form, Query
from fastapi.responses import HTMLResponse, Response, JSONResponse, StreamingResponse
from datetime import date, datetime, timedelta, timezone
from io import BytesIO
import json
import re
import random
import urllib.request
import urllib.error
import ssl
import time
from typing import Optional
from zoneinfo import ZoneInfo

from docx import Document

from app.db import migrate, connect
from app.seed import seed_all

app = FastAPI(title="Cert IV Coach (Offline)")

INDEX_HTML = r"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>Cert IV Coach</title>
  <link rel="stylesheet" href="/styles.css" />
</head>
<body>
<div class="app-shell">
  <aside class="sidebar">
    <div class="sidebar-brand">
      <div class="brand-icon">C4</div>
      <div class="brand-text">
        <span class="brand-name">Cert IV Coach</span>
        <span class="brand-sub">Offline Study Tool</span>
      </div>
    </div>

    <div class="sidebar-unit">
      <label class="sidebar-label">Active Unit</label>
      <select id="unitSelect" onchange="selectUnitFromDropdown()">
        <option value="">Choose a unit...</option>
      </select>
    </div>

    <nav class="sidebar-nav" id="sidebarNav">
      <div class="nav-group-label">Study</div>
      <button class="nav-item active" data-page="dashboard" onclick="showPage('dashboard')">
        <svg class="nav-icon" viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.5"><rect x="2" y="2" width="7" height="7" rx="1.5"/><rect x="11" y="2" width="7" height="7" rx="1.5"/><rect x="2" y="11" width="7" height="7" rx="1.5"/><rect x="11" y="11" width="7" height="7" rx="1.5"/></svg>
        Dashboard
      </button>
      <button class="nav-item" data-page="quiz" onclick="showPage('quiz')">
        <svg class="nav-icon" viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.5"><circle cx="10" cy="10" r="8"/><path d="M7.5 7.5a2.5 2.5 0 014.5 1.5c0 1.5-2.5 2-2.5 3.5"/><circle cx="10" cy="15" r="0.75" fill="currentColor" stroke="none"/></svg>
        Practice Quiz
      </button>
      <button class="nav-item" data-page="short-answer" onclick="showPage('short-answer')">
        <svg class="nav-icon" viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.5"><rect x="3" y="2" width="14" height="16" rx="2"/><line x1="6" y1="6" x2="14" y2="6"/><line x1="6" y1="9.5" x2="14" y2="9.5"/><line x1="6" y1="13" x2="11" y2="13"/></svg>
        Short Answer
      </button>
      <button class="nav-item" data-page="exam" onclick="showPage('exam')">
        <svg class="nav-icon" viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.5"><rect x="3" y="3" width="14" height="14" rx="2"/><path d="M6 7h8M6 10h5M6 13h3"/></svg>
        Exam Mode
      </button>
      <button class="nav-item" data-page="mistakes" onclick="showPage('mistakes')">
        <svg class="nav-icon" viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.5"><circle cx="6" cy="6" r="3"/><circle cx="14" cy="6" r="3"/><path d="M3 14c1.5-2 3.5-3 7-3s5.5 1 7 3"/></svg>
        Mistakes
      </button>
      <button class="nav-item" data-page="flashcards" onclick="showPage('flashcards')">
        <svg class="nav-icon" viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.5"><rect x="1" y="5" width="12" height="9" rx="1.5"/><rect x="5" y="2" width="12" height="9" rx="1.5"/></svg>
        Flashcards
      </button>
      <button class="nav-item" data-page="explain" onclick="showPage('explain')">
        <svg class="nav-icon" viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.5"><path d="M4 4h12v8H8l-4 4z"/><circle cx="9" cy="8" r="1"/><circle cx="12" cy="8" r="1"/></svg>
        Explain It Back
      </button>

      <div class="nav-group-label">Assessment</div>
      <button class="nav-item" data-page="case-study" onclick="showPage('case-study')">
        <svg class="nav-icon" viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.5"><path d="M5 2h7l4 4v11a1 1 0 01-1 1H5a1 1 0 01-1-1V3a1 1 0 011-1z"/><polyline points="12,2 12,6 16,6"/></svg>
        Case Study Builder
      </button>
      <button class="nav-item" data-page="add-card" onclick="showPage('add-card')">
        <svg class="nav-icon" viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="2"><line x1="10" y1="4" x2="10" y2="16"/><line x1="4" y1="10" x2="16" y2="10"/></svg>
        Add Card
      </button>
    </nav>

    <div class="sidebar-footer">
      <button class="nav-item" data-page="settings" onclick="showPage('settings')">
        <svg class="nav-icon" viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.5"><circle cx="10" cy="10" r="3"/><path d="M10 1.5v2.5M10 16v2.5M1.5 10H4M16 10h2.5M4.1 4.1l1.8 1.8M14.1 14.1l1.8 1.8M4.1 15.9l1.8-1.8M14.1 5.9l1.8-1.8"/></svg>
        Settings
      </button>
      <div class="sidebar-status">
        <span class="status-dot"></span>
        <span>Offline Mode</span>
      </div>
    </div>
  </aside>

  <main class="main" id="mainContent">
    <!-- Dashboard -->
    <div id="page-dashboard" class="page active">
      <div class="page-title-bar">
        <h1 id="unitTitle">Select a unit to begin</h1>
      </div>
      <div class="dashboard-welcome" id="dashboardWelcome">
        <div class="welcome-icon">
          <svg viewBox="0 0 48 48" fill="none" stroke="currentColor" stroke-width="1.5" width="48" height="48"><rect x="6" y="6" width="16" height="16" rx="3" opacity="0.4"/><rect x="26" y="6" width="16" height="16" rx="3" opacity="0.6"/><rect x="6" y="26" width="16" height="16" rx="3" opacity="0.6"/><rect x="26" y="26" width="16" height="16" rx="3" opacity="0.4"/></svg>
        </div>
        <p class="welcome-text">Choose a unit from the sidebar to start studying.</p>
      </div>
      <div id="dashboardContent" style="display:none;">
        <div class="stats-grid">
          <div class="stat-card">
            <div class="stat-label">Quiz Score</div>
            <div class="stat-value" id="dashQuizPct">--</div>
            <div class="stat-sub" id="dashQuizDetail">No attempts yet</div>
          </div>
          <div class="stat-card">
            <div class="stat-label">Short Answer</div>
            <div class="stat-value" id="dashSAPct">--</div>
            <div class="stat-sub" id="dashSADetail">No attempts yet</div>
          </div>
          <div class="stat-card">
            <div class="stat-label">Cards Due</div>
            <div class="stat-value" id="dashCardsDue">--</div>
            <div class="stat-sub">for review today</div>
          </div>
        </div>
        <div class="panel" id="dashProgressPanel">
          <h3>Progress (last 14 days)</h3>
          <div id="progressDash"></div>
        </div>
        <div class="panel" id="dashWeakTags" style="display:none;">
          <h3>Areas to Improve</h3>
          <div id="dashWeakTagsList" class="quiz-stats-tags"></div>
        </div>
      </div>
    </div>

    <!-- Quiz -->
    <div id="page-quiz" class="page">
      <div class="page-title-bar">
        <h2>Practice Quiz</h2>
        <span class="page-subtitle">Assessment 1 preparation</span>
      </div>
      <div class="content-area">
        <div class="action-bar">
          <button onclick="startQuiz()">Start / Next</button>
          <button class="btn-secondary" id="quizPrevBtn" onclick="prevQuiz()">Previous</button>
          <button class="btn-ghost" id="quizToggleContext" onclick="toggleQuizContext()">Show hint</button>
        </div>
        <div class="card" id="quizBox" style="display:none;">
          <div id="quizQ" class="question-text"></div>
          <div id="quizContext" class="hint-box" style="display:none;"><strong>Hint:</strong> <span id="quizContextText"></span></div>
          <div id="quizChoices"></div>
          <div class="action-bar" style="margin-top:12px;">
            <button id="quizSubmitBtn" onclick="submitQuiz()">Submit</button>
            <button class="btn-secondary" onclick="startQuiz()">Skip</button>
            <button class="btn-ghost" onclick="showQuizAskUI()">Ask about this question</button>
          </div>
          <div id="quizAskArea" class="challenge-area"></div>
          <div id="quizFeedback" class="feedback-area"></div>
          <div id="quizChallengeArea" class="challenge-area"></div>
        </div>
        <div id="quizStats" class="quiz-stats"></div>
      </div>
    </div>

    <!-- Short Answer -->
    <div id="page-short-answer" class="page">
      <div class="page-title-bar">
        <h2>Short Answer Practice</h2>
        <span class="page-subtitle">Written response training</span>
      </div>
      <div class="content-area">
        <div class="action-bar">
          <button onclick="startShortAnswer()">Start / Next</button>
          <button class="btn-secondary" id="saPrevBtn" onclick="prevShortAnswer()">Previous</button>
          <button class="btn-ghost" id="saToggleContext" onclick="toggleSAContext()">Show hint</button>
        </div>
        <div class="card" id="saBox" style="display:none;">
          <div id="saQ" class="question-text"></div>
          <div id="saContext" class="hint-box" style="display:none;"><strong>Hint:</strong> <span id="saContextText"></span></div>
          <textarea id="saAnswer" rows="5" placeholder="Type your answer here..."></textarea>
          <div class="action-bar" style="margin-top:12px;">
            <button onclick="submitShortAnswer()" id="saSubmitBtn">Submit</button>
            <button class="btn-teach" onclick="teachMeSA()" id="saTeachBtn">I don't know &mdash; Teach me</button>
            <button class="btn-secondary" onclick="startShortAnswer()">Skip</button>
          </div>
          <div id="saFeedback" class="feedback-area"></div>
          <div id="saChallengeArea" class="challenge-area"></div>
          <div id="saModelAnswer" style="display:none;" class="sa-model-answer">
            <strong>Model answer:</strong>
            <div id="saModelAnswerText"></div>
          </div>
        </div>
        <div id="saStats" class="quiz-stats"></div>
      </div>
    </div>

    <!-- Mistakes -->
    <div id="page-mistakes" class="page">
      <div class="page-title-bar">
        <h2>Mistakes Review</h2>
        <span class="page-subtitle">Target your weakest questions</span>
      </div>
      <div class="content-area">
        <p class="text-muted">These are questions you've previously answered incorrectly. Filter by topic and re-attempt them.</p>
        <div id="mistakeTags" class="quiz-stats-tags"></div>
        <div id="mistakeList"></div>
      </div>
    </div>

    <!-- Exam Mode -->
    <div id="page-exam" class="page">
      <div class="page-title-bar">
        <h2>Exam Mode</h2>
        <span class="page-subtitle">Timed mock assessment</span>
      </div>
      <div class="content-area" id="examContent">
        <!-- Filled by JS: setup, in-progress, or results -->
      </div>
    </div>

    <!-- Explain It Back -->
    <div id="page-explain" class="page">
      <div class="page-title-bar">
        <h2>Explain It Back</h2>
        <span class="page-subtitle">Teach the concept in your own words</span>
      </div>
      <div class="content-area" id="explainContent">
        <!-- Filled by JS -->
      </div>
    </div>

    <!-- Flashcards -->
    <div id="page-flashcards" class="page">
      <div class="page-title-bar">
        <h2>Flashcards</h2>
        <span class="page-subtitle">Spaced repetition review</span>
      </div>
      <div class="content-area">
        <div class="action-bar">
          <button class="btn-secondary" onclick="setMode('Knowledge')">Knowledge</button>
          <button class="btn-secondary" onclick="setMode('Case Study')">Case Study</button>
          <button class="btn-secondary" onclick="setMode('Project')">Project</button>
        </div>
        <p class="text-muted" style="margin-top:10px;">Cards due for review today in the selected mode.</p>
        <div id="review"></div>
      </div>
    </div>

    <!-- Case Study -->
    <div id="page-case-study" class="page">
      <div class="page-title-bar">
        <h2>Case Study Builder</h2>
        <span class="page-subtitle">Assessment 2 preparation</span>
      </div>
      <div class="content-area">
        <p class="text-muted">Pick a task, write your response, save, then export as DOCX.</p>
        <label class="field-label">AE2 item</label>
        <select id="ae2ItemSelect" onchange="loadAE2Item()"></select>
        <label class="field-label">Response <span class="text-muted">(supports markdown tables with | pipes)</span></label>
        <textarea id="ae2Content" rows="14"></textarea>
        <div class="action-bar">
          <button onclick="saveAE2()">Save</button>
          <button class="btn-secondary" onclick="exportAE2()">Export DOCX</button>
        </div>
        <div class="text-muted" id="ae2Status"></div>
      </div>
    </div>

    <!-- Add Card -->
    <div id="page-add-card" class="page">
      <div class="page-title-bar">
        <h2>Add a Card</h2>
        <span class="page-subtitle">Create custom flashcards</span>
      </div>
      <div class="content-area">
        <form id="cardForm">
          <input type="hidden" name="unit_id" id="card_unit_id" />
          <input type="hidden" name="event_kind" id="card_event_kind" value="Knowledge" />
          <label class="field-label">Prompt</label>
          <textarea name="prompt" rows="2" required></textarea>
          <label class="field-label">Answer</label>
          <textarea name="answer" rows="2" required></textarea>
          <label class="field-label">Tags (comma-separated)</label>
          <input name="tags" />
          <button type="submit">Add card</button>
        </form>
      </div>
    </div>

    <!-- Settings -->
    <div id="page-settings" class="page">
      <div class="page-title-bar">
        <h2>Settings</h2>
        <span class="page-subtitle">Configure your study environment</span>
      </div>
      <div class="content-area">
        <div class="settings-section">
          <h3>Student Details</h3>
          <label class="field-label">Name</label>
          <input id="student_name" />
          <label class="field-label">Student number (optional)</label>
          <input id="student_number" />
          <button onclick="saveSettings()">Save</button>
          <div class="text-muted" id="settingsSaved"></div>
        </div>

        <div class="settings-section">
          <h3>Quiz Behaviour</h3>
          <div class="setting-row">
            <div class="setting-info">
              <div class="setting-title">Require submit before next</div>
              <div class="text-muted">Prevents skipping ahead until you answer.</div>
            </div>
            <label class="switch">
              <input type="checkbox" id="settings_require_submit" />
              <span class="switch-track"><span class="switch-thumb"></span></span>
            </label>
          </div>
        </div>

        <div class="settings-section">
          <h3>AI Grading</h3>
          <label class="field-label">Gemini API key <span class="text-muted">(for AI grading of short answers)</span></label>
          <input type="password" id="settings_gemini_key" placeholder="Paste your API key from ai.google.dev" />
          <div class="action-bar" style="margin-top:6px;">
            <button onclick="saveGeminiKey()">Save key</button>
            <span class="text-muted" id="geminiKeySaved" style="align-self:center;"></span>
          </div>
          <p class="text-muted">Free at <a href="https://aistudio.google.com/apikey" target="_blank" class="link">ai.google.dev</a>. Without a key, short answers use self-grading.</p>
        </div>
      </div>
    </div>

  </main>
</div>

<!-- Glossary modal -->
<div id="glossaryModal" class="modal" style="display:none;">
  <div class="modal-overlay" onclick="closeGlossary()"></div>
  <div class="modal-content">
    <div class="modal-header">
      <h3 id="glossaryTitle">Term</h3>
      <button class="btn-ghost modal-close" onclick="closeGlossary()">Close</button>
    </div>
    <div id="glossaryBody" class="modal-body"></div>
  </div>
</div>

<script src="/app.js"></script>
</body>
</html>
"""

STYLES_CSS = r"""
:root {
  --bg-body: #080a14;
  --bg-sidebar: #0c0f1a;
  --bg-elevated: rgba(20, 23, 34, 0.95);
  --bg-card: rgba(16, 19, 28, 0.92);
  --bg-input: rgba(13, 16, 24, 0.9);
  --border-subtle: rgba(148, 163, 184, 0.1);
  --border-medium: rgba(148, 163, 184, 0.18);
  --border-strong: rgba(148, 163, 184, 0.3);
  --accent: #fb923c;
  --accent-soft: rgba(251, 146, 60, 0.1);
  --accent-medium: rgba(251, 146, 60, 0.22);
  --accent-strong: rgba(251, 146, 60, 0.45);
  --accent-secondary: #38bdf8;
  --accent-secondary-soft: rgba(56, 189, 248, 0.1);
  --text-main: #f1f5f9;
  --text-secondary: #cbd5e1;
  --text-muted: #64748b;
  --success: #34d399;
  --success-soft: rgba(52, 211, 153, 0.12);
  --warning: #fbbf24;
  --warning-soft: rgba(251, 191, 36, 0.12);
  --danger: #f87171;
  --danger-soft: rgba(248, 113, 113, 0.12);
  --radius-xs: 6px;
  --radius-sm: 8px;
  --radius-md: 12px;
  --radius-lg: 16px;
  --radius-xl: 20px;
  --shadow: 0 4px 24px rgba(0, 0, 0, 0.35);
  --shadow-lg: 0 8px 40px rgba(0, 0, 0, 0.45);
  --transition: 0.2s cubic-bezier(0.4, 0, 0.2, 1);
}

*, *::before, *::after { box-sizing: border-box; }

body {
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", system-ui, sans-serif;
  margin: 0;
  background: var(--bg-body);
  color: var(--text-main);
  -webkit-font-smoothing: antialiased;
  line-height: 1.5;
  user-select: text;
}

/* ── App Shell ── */
.app-shell {
  display: grid;
  grid-template-columns: 260px minmax(0, 1fr);
  height: 100vh;
}

/* ── Sidebar ── */
.sidebar {
  background: var(--bg-sidebar);
  border-right: 1px solid var(--border-subtle);
  display: flex;
  flex-direction: column;
  overflow: hidden;
}

.sidebar-brand {
  display: flex;
  align-items: center;
  gap: 10px;
  padding: 18px 16px 14px;
  border-bottom: 1px solid var(--border-subtle);
}

.brand-icon {
  width: 36px;
  height: 36px;
  border-radius: var(--radius-sm);
  background: linear-gradient(135deg, var(--accent), #f97316);
  display: flex;
  align-items: center;
  justify-content: center;
  font-weight: 800;
  font-size: 0.75rem;
  color: #111;
  flex-shrink: 0;
}

.brand-name {
  font-weight: 700;
  font-size: 0.95rem;
  display: block;
}

.brand-sub {
  font-size: 0.72rem;
  color: var(--text-muted);
  display: block;
}

/* Unit selector */
.sidebar-unit {
  padding: 14px 12px 6px;
}

.sidebar-label {
  font-size: 0.68rem;
  font-weight: 600;
  text-transform: uppercase;
  letter-spacing: 0.08em;
  color: var(--text-muted);
  display: block;
  margin-bottom: 4px;
  padding-left: 2px;
}

.sidebar-unit select {
  width: 100%;
  margin: 0;
  padding: 8px 10px;
  border-radius: var(--radius-sm);
  border: 1px solid var(--border-medium);
  background: var(--bg-input);
  color: var(--text-main);
  font: inherit;
  font-size: 0.85rem;
  outline: none;
  cursor: pointer;
  transition: border-color var(--transition);
}

.sidebar-unit select:focus {
  border-color: var(--accent-strong);
}

/* Navigation */
.sidebar-nav {
  flex: 1;
  overflow-y: auto;
  padding: 6px 8px;
}

.nav-group-label {
  font-size: 0.66rem;
  font-weight: 600;
  text-transform: uppercase;
  letter-spacing: 0.1em;
  color: var(--text-muted);
  padding: 14px 10px 4px;
}

.nav-item {
  width: 100%;
  display: flex;
  align-items: center;
  gap: 10px;
  padding: 9px 12px;
  margin: 1px 0;
  border: none;
  border-radius: var(--radius-sm);
  background: transparent;
  color: var(--text-secondary);
  font: inherit;
  font-size: 0.88rem;
  cursor: pointer;
  text-align: left;
  transition: all var(--transition);
  box-shadow: none;
}

.nav-item:hover {
  background: rgba(148, 163, 184, 0.06);
  color: var(--text-main);
  transform: none;
  box-shadow: none;
}

.nav-item.active {
  background: var(--accent-soft);
  color: var(--accent);
  font-weight: 500;
}

.nav-icon {
  width: 18px;
  height: 18px;
  flex-shrink: 0;
  opacity: 0.65;
}

.nav-item.active .nav-icon {
  opacity: 1;
}

/* Sidebar footer */
.sidebar-footer {
  padding: 6px 8px 8px;
  border-top: 1px solid var(--border-subtle);
}

.sidebar-status {
  display: flex;
  align-items: center;
  gap: 8px;
  padding: 8px 12px;
  font-size: 0.73rem;
  color: var(--text-muted);
}

.status-dot {
  width: 7px;
  height: 7px;
  border-radius: 50%;
  background: var(--success);
  box-shadow: 0 0 8px rgba(52, 211, 153, 0.4);
}

/* ── Main Content ── */
.main {
  overflow-y: auto;
  background: var(--bg-body);
}

.page {
  display: none;
  padding: 28px 32px 40px;
  max-width: 820px;
  animation: pageFadeIn 0.18s ease;
}

.page.active {
  display: block;
}

@keyframes pageFadeIn {
  from { opacity: 0; transform: translateY(6px); }
  to { opacity: 1; transform: translateY(0); }
}

.page-title-bar {
  margin-bottom: 22px;
  padding-bottom: 14px;
  border-bottom: 1px solid var(--border-subtle);
}

.page-title-bar h1,
.page-title-bar h2 {
  margin: 0;
  font-size: 1.35rem;
  font-weight: 700;
  letter-spacing: -0.01em;
}

.page-subtitle {
  font-size: 0.82rem;
  color: var(--text-muted);
  margin-top: 2px;
  display: block;
}

/* ── Dashboard ── */
.dashboard-welcome {
  padding: 48px 0;
  text-align: center;
}

.welcome-icon {
  color: var(--accent);
  opacity: 0.5;
  margin-bottom: 12px;
}

.welcome-text {
  color: var(--text-muted);
  font-size: 0.95rem;
}

.stats-grid {
  display: grid;
  grid-template-columns: repeat(3, 1fr);
  gap: 14px;
  margin-bottom: 20px;
}

.stat-card {
  background: var(--bg-elevated);
  border: 1px solid var(--border-subtle);
  border-radius: var(--radius-lg);
  padding: 18px 20px;
}

.stat-label {
  font-size: 0.75rem;
  color: var(--text-muted);
  font-weight: 500;
  text-transform: uppercase;
  letter-spacing: 0.06em;
}

.stat-value {
  font-size: 2rem;
  font-weight: 700;
  color: var(--accent);
  margin: 4px 0 2px;
  line-height: 1.1;
}

.stat-sub {
  font-size: 0.78rem;
  color: var(--text-muted);
}

/* ── Panels & Cards ── */
.panel {
  background: var(--bg-elevated);
  border: 1px solid var(--border-subtle);
  border-radius: var(--radius-lg);
  padding: 16px 18px;
  margin: 16px 0;
}

.panel h3 {
  margin: 0 0 10px;
  font-size: 1rem;
  font-weight: 600;
}

.card {
  border: 1px solid var(--border-medium);
  border-radius: var(--radius-lg);
  padding: 14px 16px;
  margin: 14px 0;
  background: var(--bg-card);
}

/* ── Buttons ── */
.action-bar {
  display: flex;
  gap: 8px;
  flex-wrap: wrap;
  align-items: center;
}

.row {
  display: flex;
  gap: 10px;
  flex-wrap: wrap;
}

button {
  padding: 8px 16px;
  border-radius: var(--radius-sm);
  border: none;
  background: linear-gradient(135deg, #fb923c, #f97316);
  color: #111;
  cursor: pointer;
  font: inherit;
  font-size: 0.88rem;
  font-weight: 600;
  transition: all var(--transition);
  box-shadow: none;
}

button:hover {
  filter: brightness(1.1);
  transform: translateY(-1px);
  box-shadow: 0 4px 14px rgba(251, 146, 60, 0.25);
}

button:active {
  transform: translateY(0);
  box-shadow: none;
}

button:disabled {
  opacity: 0.4;
  cursor: default;
  transform: none;
  filter: none;
  box-shadow: none;
}

.btn-secondary {
  background: var(--bg-elevated);
  color: var(--text-main);
  border: 1px solid var(--border-medium);
}

.btn-secondary:hover {
  background: rgba(148, 163, 184, 0.08);
  border-color: var(--border-strong);
  filter: none;
  box-shadow: none;
}

.btn-ghost {
  background: transparent;
  color: var(--text-secondary);
  border: 1px solid transparent;
}

.btn-ghost:hover {
  background: rgba(148, 163, 184, 0.06);
  filter: none;
  box-shadow: none;
}

.btn-teach {
  background: linear-gradient(135deg, #38bdf8, #0ea5e9);
  color: #111;
  border: none;
}

.btn-teach:hover {
  filter: brightness(1.1);
  box-shadow: 0 4px 14px rgba(56, 189, 248, 0.25);
}

.modal-close {
  background: var(--bg-elevated);
  border: 1px solid var(--border-medium);
  color: var(--text-main);
  padding: 6px 14px;
  border-radius: var(--radius-sm);
  font-size: 0.82rem;
}

/* ── Forms ── */
textarea,
input:not([type="radio"]):not([type="checkbox"]),
select {
  width: 100%;
  margin: 4px 0 12px;
  padding: 9px 12px;
  border-radius: var(--radius-sm);
  border: 1px solid var(--border-medium);
  background: var(--bg-input);
  color: var(--text-main);
  font: inherit;
  font-size: 0.9rem;
  outline: none;
  transition: border-color var(--transition), box-shadow var(--transition);
}

textarea:focus,
input:not([type="radio"]):not([type="checkbox"]):focus,
select:focus {
  border-color: var(--accent-strong);
  box-shadow: 0 0 0 2px var(--accent-soft);
}

input[type="radio"] {
  appearance: none;
  -webkit-appearance: none;
  width: 18px;
  height: 18px;
  min-width: 18px;
  border-radius: 50%;
  border: 2px solid var(--border-strong);
  background: transparent;
  margin: 0;
  padding: 0;
  cursor: pointer;
  transition: all var(--transition);
}

input[type="radio"]:checked {
  border-color: var(--accent);
  background: var(--accent);
  box-shadow: inset 0 0 0 3px var(--bg-card);
}

.field-label {
  font-size: 0.85rem;
  font-weight: 500;
  color: var(--text-secondary);
  display: block;
  margin-top: 8px;
}

/* ── Text utilities ── */
.text-muted,
.muted {
  font-size: 0.84rem;
  color: var(--text-muted);
}

.link {
  color: var(--accent-secondary);
  text-decoration: none;
}

.link:hover {
  text-decoration: underline;
}

/* ── Question styling ── */
.question-text {
  font-weight: 600;
  font-size: 1rem;
  margin-bottom: 10px;
  line-height: 1.5;
}

.hint-box {
  margin: 8px 0;
  padding: 10px 12px;
  background: var(--accent-secondary-soft);
  border: 1px solid rgba(56, 189, 248, 0.18);
  border-radius: var(--radius-sm);
  font-size: 0.88rem;
  color: var(--text-secondary);
}

.feedback-area {
  margin-top: 12px;
}

.challenge-area {
  margin-top: 8px;
}

.challenge-box {
  border: 1px solid rgba(15, 118, 110, 0.2);
  background: var(--accent-secondary-soft);
  border-radius: var(--radius-sm);
  padding: 10px;
}

.challenge-header {
  font-size: 0.85rem;
  font-weight: 600;
  margin-bottom: 6px;
}

.challenge-textarea {
  width: 100%;
  border-radius: var(--radius-sm);
  border: 1px solid rgba(148, 163, 184, 0.6);
  padding: 8px;
  min-height: 70px;
  resize: vertical;
  font-family: inherit;
  font-size: 0.92rem;
}

.challenge-actions {
  margin-top: 6px;
}

.challenge-result {
  margin-top: 8px;
  font-size: 0.9rem;
}

.challenge-status,
.challenge-error {
  color: var(--danger);
}

.challenge-clarification {
  color: var(--text-main);
}

.challenge-trigger {
  padding: 0;
  border: none;
  background: none;
  color: var(--accent);
  font-size: 0.9rem;
  font-weight: 600;
  cursor: pointer;
}
/* ── Quiz choices ── */
#quizChoices {
  display: flex;
  flex-direction: column;
  gap: 6px;
}

.quiz-choice-row {
  margin-bottom: 0;
  width: 100%;
}

.quiz-choice {
  display: flex;
  align-items: center;
  gap: 10px;
  line-height: 1.4;
  cursor: pointer;
  width: 100%;
  padding: 9px 12px;
  border-radius: var(--radius-sm);
  transition: background var(--transition);
}

.quiz-choice:hover {
  background: rgba(148, 163, 184, 0.05);
}

.quiz-choice input {
  flex-shrink: 0;
}

.quiz-choice-text {
  display: block;
  text-align: left;
  flex: 1;
}

.quiz-counter {
  font-size: 0.78rem;
  color: var(--text-muted);
  margin-bottom: 6px;
  letter-spacing: 0.04em;
}

/* ── Quiz stats ── */
.quiz-stats {
  margin-top: 16px;
}

.quiz-stats-header {
  display: flex;
  flex-direction: column;
  gap: 8px;
}

.quiz-stats-score {
  display: flex;
  align-items: baseline;
  gap: 8px;
}

.quiz-stats-big {
  font-size: 1.5rem;
  font-weight: 700;
  color: var(--accent);
}

.quiz-stats-detail {
  font-size: 0.82rem;
  color: var(--text-muted);
}

.quiz-stats-bar-track {
  width: 100%;
  height: 6px;
  background: rgba(148, 163, 184, 0.1);
  border-radius: 999px;
  overflow: hidden;
}

.quiz-stats-bar-fill {
  height: 100%;
  background: linear-gradient(90deg, var(--accent), #fdba74);
  border-radius: 999px;
  transition: width 0.4s ease;
}

.quiz-stats-tags {
  display: flex;
  flex-wrap: wrap;
  gap: 6px;
  margin-top: 10px;
}

.tag-chip {
  padding: 3px 9px;
  border-radius: var(--radius-xs);
  font-size: 0.76rem;
  border: 1px solid var(--border-soft);
  background: var(--bg-elevated);
  color: var(--text-secondary);
  cursor: pointer;
}

.tag-chip-active {
  border-color: var(--accent-medium);
  background: var(--accent-soft);
  color: var(--accent);
}

.mistake-card {
  margin-top: 10px;
}

.mistake-header {
  display: flex;
  justify-content: space-between;
  align-items: center;
  margin-bottom: 6px;
}

.mistake-meta {
  display: flex;
  flex-wrap: wrap;
  gap: 6px;
  align-items: center;
  font-size: 0.78rem;
}

.mistake-badge {
  padding: 2px 7px;
  border-radius: 999px;
  font-size: 0.72rem;
  text-transform: uppercase;
  letter-spacing: 0.08em;
}

.mistake-badge.mastered {
  background: rgba(34, 197, 94, 0.12);
  color: #22c55e;
}

.mistake-badge.pending {
  background: rgba(248, 113, 113, 0.12);
  color: #f97373;
}

.mistake-count {
  padding: 0 6px;
  border-left: 1px solid var(--border-soft);
  font-size: 0.74rem;
}

.mistake-tags {
  font-size: 0.74rem;
  color: var(--text-muted);
}

.mistake-source {
  padding: 0 6px;
  border-left: 1px solid var(--border-soft);
  font-size: 0.74rem;
  font-weight: 500;
  color: var(--accent);
}

.mistake-question {
  font-weight: 500;
  margin-bottom: 6px;
}

.mistake-detail {
  font-size: 0.85rem;
  color: var(--text-main);
}

.mistake-actions {
  margin-top: 8px;
}

.exam-header {
  display: flex;
  justify-content: space-between;
  align-items: center;
  margin-bottom: 8px;
  font-size: 0.9rem;
}

.exam-timer {
  font-weight: 600;
  color: var(--accent);
}

.exam-card {
  margin-top: 4px;
}

.exam-review-item {
  margin-top: 8px;
  padding-top: 6px;
  border-top: 1px solid var(--border-soft);
  font-size: 0.9rem;
}

.exam-review-header {
  display: flex;
  align-items: center;
  gap: 6px;
}

.exam-review-body {
  margin-top: 6px;
}

.exam-review-actions {
  margin-top: 6px;
}

.badge-correct,
.badge-incorrect {
  padding: 2px 7px;
  border-radius: 999px;
  font-size: 0.72rem;
}

.badge-correct {
  background: rgba(34, 197, 94, 0.12);
  color: #22c55e;
}

.badge-incorrect {
  background: rgba(248, 113, 113, 0.12);
  color: #f97373;
}

.quiz-stats-tag {
  padding: 4px 10px;
  border-radius: var(--radius-xs);
  font-size: 0.78rem;
  background: var(--accent-soft);
  border: 1px solid var(--accent-medium);
  color: var(--text-main);
}

.quiz-stats-tag strong {
  color: var(--accent);
}

/* ── Toggle switch ── */
.switch {
  display: inline-flex;
  align-items: center;
  cursor: pointer;
  user-select: none;
}

.switch input { display: none; }

.switch-track {
  width: 42px;
  height: 24px;
  background: rgba(148, 163, 184, 0.22);
  border-radius: 999px;
  position: relative;
  transition: background var(--transition);
}

.switch-thumb {
  width: 18px;
  height: 18px;
  background: #f8fafc;
  border-radius: 50%;
  position: absolute;
  top: 3px;
  left: 3px;
  transition: transform var(--transition);
  box-shadow: 0 2px 6px rgba(0, 0, 0, 0.3);
}

.switch input:checked + .switch-track {
  background: var(--accent-medium);
}

.switch input:checked + .switch-track .switch-thumb {
  transform: translateX(18px);
}

/* ── Settings ── */
.settings-section {
  padding: 18px 0;
  border-bottom: 1px solid var(--border-subtle);
}

.settings-section:last-child {
  border-bottom: none;
}

.settings-section h3 {
  margin: 0 0 12px;
  font-size: 1rem;
  font-weight: 600;
}

.setting-row {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 16px;
}

.setting-info { flex: 1; }

.setting-title {
  font-weight: 500;
  font-size: 0.9rem;
}

/* ── Short answer feedback ── */
.sa-verdict {
  display: inline-block;
  padding: 3px 12px;
  border-radius: var(--radius-xs);
  font-size: 0.82rem;
  font-weight: 600;
  margin-bottom: 6px;
}

.sa-verdict-correct {
  background: var(--success-soft);
  border: 1px solid rgba(52, 211, 153, 0.35);
  color: #6ee7b7;
}

.sa-verdict-partial {
  background: var(--warning-soft);
  border: 1px solid rgba(251, 191, 36, 0.35);
  color: #fde68a;
}

.sa-verdict-incorrect {
  background: var(--danger-soft);
  border: 1px solid rgba(248, 113, 113, 0.35);
  color: #fca5a5;
}

.sa-feedback-text {
  font-size: 0.88rem;
  color: var(--text-muted);
  margin-top: 4px;
  line-height: 1.5;
}

.sa-model-answer {
  background: rgba(148, 163, 184, 0.04);
  border: 1px solid var(--border-subtle);
  border-radius: var(--radius-sm);
  padding: 12px 14px;
  margin-top: 12px;
  font-size: 0.88rem;
  line-height: 1.5;
}

.sa-self-grade {
  margin-top: 8px;
}

.sa-self-grade button {
  font-size: 0.82rem;
  padding: 5px 12px;
}

.sa-grading {
  color: var(--text-muted);
  font-size: 0.86rem;
  font-style: italic;
}

/* ── Teach me box ── */
.sa-teach-box {
  background: rgba(56, 189, 248, 0.06);
  border: 1px solid rgba(56, 189, 248, 0.2);
  border-radius: var(--radius-sm);
  padding: 12px 14px;
  margin-top: 8px;
  font-size: 0.88rem;
  line-height: 1.6;
  color: var(--text-main);
}

.sa-teach-label {
  font-size: 0.74rem;
  font-weight: 600;
  color: var(--accent-secondary);
  text-transform: uppercase;
  letter-spacing: 0.06em;
  margin-bottom: 4px;
}

.card-created-note {
  font-size: 0.8rem;
  color: var(--accent-secondary);
  margin-top: 6px;
  font-style: italic;
}

/* ── Progress dashboard ── */
.progress-header {
  display: flex;
  gap: 14px;
  flex-wrap: wrap;
  margin-bottom: 14px;
}

.progress-stat {
  display: flex;
  flex-direction: column;
  align-items: center;
  padding: 10px 16px;
  border-radius: var(--radius-sm);
  background: rgba(148, 163, 184, 0.04);
  border: 1px solid var(--border-subtle);
  min-width: 80px;
}

.progress-stat-big {
  font-size: 1.3rem;
  font-weight: 700;
  color: var(--accent);
}

.progress-stat-label {
  font-size: 0.72rem;
  color: var(--text-muted);
  text-transform: uppercase;
  letter-spacing: 0.06em;
  margin-top: 2px;
}

.progress-chart {
  display: flex;
  align-items: flex-end;
  gap: 4px;
  height: 80px;
  padding: 4px 0;
}

.progress-bar {
  flex: 1;
  min-width: 0;
  border-radius: 4px 4px 0 0;
  position: relative;
  transition: height 0.3s ease;
}

.progress-bar-fill {
  width: 100%;
  border-radius: 4px 4px 0 0;
  position: absolute;
  bottom: 0;
  background: linear-gradient(180deg, var(--accent), rgba(251, 146, 60, 0.4));
  transition: height 0.3s ease;
}

.progress-bar-label {
  position: absolute;
  bottom: -16px;
  left: 50%;
  transform: translateX(-50%);
  font-size: 0.6rem;
  color: var(--text-muted);
  white-space: nowrap;
}

.progress-chart-wrap {
  margin-top: 4px;
  padding-bottom: 18px;
}

/* ── Glossary modal ── */
.modal {
  position: fixed;
  inset: 0;
  display: none;
  z-index: 9999;
}

.modal-overlay {
  position: absolute;
  inset: 0;
  background: rgba(0, 0, 0, 0.55);
  backdrop-filter: blur(4px);
}

.modal-content {
  position: relative;
  width: 480px;
  max-width: 90%;
  margin: 8% auto;
  background: var(--bg-elevated);
  border: 1px solid var(--border-medium);
  padding: 18px 20px;
  border-radius: var(--radius-xl);
  color: var(--text-main);
  box-shadow: var(--shadow-lg);
}

.modal-header {
  display: flex;
  justify-content: space-between;
  align-items: center;
}

.modal-header h3 { margin: 0; }

.modal-body {
  margin-top: 10px;
  white-space: pre-wrap;
  color: var(--text-secondary);
  line-height: 1.6;
}

.glossary-term {
  color: var(--accent-secondary);
  text-decoration: underline;
  text-decoration-style: dotted;
  cursor: pointer;
}

.glossary-term:hover {
  text-decoration-style: solid;
}

/* ── Responsive ── */
@media (max-width: 768px) {
  .app-shell {
    grid-template-columns: 1fr;
  }
  .sidebar {
    display: none;
  }
  .page {
    padding: 20px 16px 32px;
  }
  .stats-grid {
    grid-template-columns: 1fr;
  }
}
"""

APP_JS = r"""
let currentUnitId = null;
let currentMode = 'Knowledge';
let currentQuiz = null;
let quizHintVisible = false;
let lastQuizSubmitted = false;
let requireSubmitBeforeNext = false;
let ae2AutosaveTimer = null;
let quizHistory = [];
let quizHistoryIndex = -1;
let quizQuestionNumber = 0;
let quizTotalQuestions = 0;
let availableUnits = [];
let unitsLoadedPromise = null;

// Short answer state
let currentSA = null;
let saHintVisible = false;
let lastSASubmitted = false;
let saHistory = [];
let saHistoryIndex = -1;
let saQuestionNumber = 0;
let saTotalQuestions = 0;
let mistakesData = null;
let mistakeFilterTag = localStorage.getItem('mistakeFilterTag') || '';
let currentExam = null;
let currentExplainTopic = null;

function shuffleArray(arr) {
  for (let i = arr.length - 1; i > 0; i--) {
    const j = Math.floor(Math.random() * (i + 1));
    [arr[i], arr[j]] = [arr[j], arr[i]];
  }
  return arr;
}

function getShuffledIndices(questionObj) {
  if (!questionObj || !Array.isArray(questionObj.choices)) return [];
  const count = questionObj.choices.length;
  if (!questionObj.shuffled_indices || questionObj.shuffled_indices.length !== count) {
    questionObj.shuffled_indices = shuffleArray(Array.from({ length: count }, (_, idx) => idx));
  }
  return questionObj.shuffled_indices;
}

// Force default: hints hidden unless user clicks Show hint
if (!localStorage.getItem('quizShowContext')) {
  localStorage.setItem('quizShowContext', 'false');
}

// ── Navigation ─────────────────────────────────────
function showPage(pageId) {
  document.querySelectorAll('.page').forEach(p => p.classList.remove('active'));
  document.querySelectorAll('.nav-item').forEach(n => n.classList.remove('active'));

  const page = document.getElementById('page-' + pageId);
  if (page) page.classList.add('active');

  const navBtn = document.querySelector('.nav-item[data-page="' + pageId + '"]');
  if (navBtn) navBtn.classList.add('active');

  try {
    localStorage.setItem('selectedPage', pageId);
  } catch (e) {
    console.error('Failed to persist selected page:', e);
  }

  if (pageId === 'dashboard' && currentUnitId) refreshDashboard();
  if (pageId === 'flashcards' && currentUnitId) loadReview();
  if (pageId === 'mistakes' && currentUnitId) loadMistakes();
  if (pageId === 'exam' && currentUnitId) showExamSetup();
  if (pageId === 'explain' && currentUnitId) loadExplainTopic();
}

// ── API helper ─────────────────────────────────────
async function api(path, options={}) {
  const res = await fetch(path, options);
  if (!res.ok) {
    const txt = await res.text();
    throw new Error(txt || `HTTP ${res.status}`);
  }
  const ct = res.headers.get('content-type') || '';
  if (ct.includes('application/json')) return await res.json();
  return await res.text();
}

function toggleButtonLoading(btn, isLoading, loadingText) {
  if (!btn) return;
  if (isLoading) {
    if (!btn.dataset.originalLabel) {
      btn.dataset.originalLabel = btn.textContent;
    }
    btn.textContent = loadingText;
    btn.disabled = true;
    return;
  }
  btn.disabled = false;
  if (btn.dataset.originalLabel) {
    btn.textContent = btn.dataset.originalLabel;
    delete btn.dataset.originalLabel;
  }
}

// ── Units ──────────────────────────────────────────
async function loadUnits() {
  const data = await api('/api/units');
  availableUnits = data.units || [];
  const sel = document.getElementById('unitSelect');
  sel.innerHTML = '<option value="">Choose a unit...</option>';
  availableUnits.forEach(u => {
    const opt = document.createElement('option');
    opt.value = JSON.stringify(u);
    opt.textContent = u.code + ' \u2014 ' + u.title;
    sel.appendChild(opt);
  });

  if (availableUnits.length === 1 && !localStorage.getItem('currentUnitId')) {
    await selectUnit(availableUnits[0]);
  }

  return data;
}

function selectUnitFromDropdown() {
  const sel = document.getElementById('unitSelect');
  if (!sel.value) return;
  const u = JSON.parse(sel.value);
  selectUnit(u);
}

async function selectUnit(u) {
  currentUnitId = u.id;
  const unitSelect = document.getElementById('unitSelect');
  if (unitSelect) {
    const serialized = JSON.stringify(u);
    if (unitSelect.value !== serialized) {
      unitSelect.value = serialized;
    }
  }
  document.getElementById('unitTitle').textContent = u.code + ' \u2014 ' + u.title;
  document.getElementById('card_unit_id').value = u.id;
  await refreshDashboard();
  await loadAE2Items();
  await loadSettings();
  mistakesData = null;
  currentExam = null;
  currentExplainTopic = null;
  // Reset quiz and short-answer so questions match the new unit
  quizHistory = [];
  quizHistoryIndex = -1;
  currentQuiz = null;
  quizQuestionNumber = 0;
  lastQuizSubmitted = false;
  saHistory = [];
  saHistoryIndex = -1;
  currentSA = null;
  saQuestionNumber = 0;
  lastSASubmitted = false;
  const quizBox = document.getElementById('quizBox');
  if (quizBox) quizBox.style.display = 'none';
  const saBox = document.getElementById('saBox');
  if (saBox) saBox.style.display = 'none';
  await refreshQuizStats();
  await refreshSAStats();
  try {
    localStorage.setItem('currentUnitId', u.id);
  } catch (e) {
    console.error('Failed to persist current unit:', e);
  }
}

// ── Dashboard ──────────────────────────────────────
async function refreshDashboard() {
  if (!currentUnitId) return;

  document.getElementById('dashboardWelcome').style.display = 'none';
  document.getElementById('dashboardContent').style.display = 'block';

  try {
    const qs = await api('/api/quiz/stats?unit_id=' + currentUnitId + '&window=20&tag_window=50');
    document.getElementById('dashQuizPct').textContent = qs.total ? qs.pct + '%' : '--';
    document.getElementById('dashQuizDetail').textContent = qs.total ? qs.correct + '/' + qs.total + ' correct' : 'No attempts yet';

    const ss = await api('/api/short-answer/stats?unit_id=' + currentUnitId + '&window=20');
    document.getElementById('dashSAPct').textContent = ss.total ? ss.pct + '%' : '--';
    document.getElementById('dashSADetail').textContent = ss.total ? ss.correct + '/' + ss.total + ' correct' : 'No attempts yet';

    const rv = await api('/api/review/today?unit_id=' + currentUnitId + '&event_kind=' + encodeURIComponent(currentMode));
    document.getElementById('dashCardsDue').textContent = rv.cards.length;

    // Weak tags
    const allTags = [].concat(qs.by_tag || [], ss.by_tag || []);
    const weakTags = allTags.filter(function(t) { return t.pct < 80; }).slice(0, 8);
    const weakEl = document.getElementById('dashWeakTags');
    const weakList = document.getElementById('dashWeakTagsList');
    if (weakTags.length > 0) {
      weakEl.style.display = 'block';
      weakList.innerHTML = weakTags.map(function(t) {
        return '<span class="quiz-stats-tag">' + t.tag + ' <strong>' + t.pct + '%</strong></span>';
      }).join('');
    } else {
      weakEl.style.display = 'none';
    }
  } catch (e) {
    console.error('Dashboard refresh error:', e);
  }

  await loadProgress();
}

// ── Quiz Stats ─────────────────────────────────────
async function refreshQuizStats() {
  if (!currentUnitId) return;

  const el = document.getElementById('quizStats');
  if (!el) return;

  const s = await api('/api/quiz/stats?unit_id=' + currentUnitId + '&window=20&tag_window=50');

  if (!s.total) {
    el.innerHTML = '';
    return;
  }

  let html = '<div class="quiz-stats-header">' +
    '<div class="quiz-stats-score">' +
    '<span class="quiz-stats-big">' + s.pct + '%</span>' +
    '<span class="quiz-stats-detail">' + s.correct + '/' + s.total + ' correct (last ' + s.window + ')</span>' +
    '</div>' +
    '<div class="quiz-stats-bar-track">' +
    '<div class="quiz-stats-bar-fill" style="width:' + Math.min(s.pct, 100) + '%"></div>' +
    '</div></div>';

  if (s.by_tag && s.by_tag.length) {
    const weakTags = s.by_tag.filter(function(t) { return t.pct < 100; }).slice(0, 5);
    if (weakTags.length) {
      html += '<div class="quiz-stats-tags">';
      weakTags.forEach(function(t) {
        html += '<span class="quiz-stats-tag">' + t.tag + ' <strong>' + t.pct + '%</strong></span>';
      });
      html += '</div>';
    }
  }

  el.innerHTML = html;
}

// ── Flashcard Mode ─────────────────────────────────
function setMode(mode) {
  currentMode = mode;
  document.getElementById('card_event_kind').value = mode;
  try {
    localStorage.setItem('flashcardMode', mode);
  } catch (e) {
    console.error('Failed to persist flashcard mode:', e);
  }
  loadReview();
}

// ── Review ─────────────────────────────────────────
async function loadReview() {
  const box = document.getElementById('review');
  box.innerHTML = '';
  if (!currentUnitId) return;

  const data = await api('/api/review/today?unit_id=' + currentUnitId + '&event_kind=' + encodeURIComponent(currentMode));
  if (data.cards.length === 0) {
    box.textContent = 'No cards due for this mode. Switch mode or add cards.';
    return;
  }

  data.cards.forEach(function(card) {
    const div = document.createElement('div');
    div.className = 'card';

    const p = document.createElement('div');
    p.textContent = card.prompt;

    const a = document.createElement('details');
    const sum = document.createElement('summary');
    sum.textContent = 'Show answer';
    const ans = document.createElement('div');
    ans.textContent = card.answer;
    a.appendChild(sum);
    a.appendChild(ans);

    const row = document.createElement('div');
    row.className = 'row';

    var grades = [
      {g:0, t:'Again'},
      {g:1, t:'Hard'},
      {g:2, t:'Good'},
      {g:3, t:'Easy'}
    ];

    grades.forEach(function(x) {
      const btn = document.createElement('button');
      btn.textContent = x.t;
      btn.onclick = async function() {
        const fd = new FormData();
        fd.append('card_id', card.id);
        fd.append('grade', x.g);
        await api('/api/review/grade', { method:'POST', body: fd });
        await loadReview();
      };
      row.appendChild(btn);
    });

    div.appendChild(p);
    div.appendChild(a);
    div.appendChild(row);
    box.appendChild(div);
  });
}

// ── Card Form ──────────────────────────────────────
document.getElementById('cardForm').addEventListener('submit', async function(ev) {
  ev.preventDefault();
  if (!currentUnitId) return alert('Select a unit first.');

  const fd = new FormData(ev.target);
  await api('/api/cards', { method:'POST', body: fd });
  ev.target.reset();
  document.getElementById('card_unit_id').value = currentUnitId;
  document.getElementById('card_event_kind').value = currentMode;
  await loadReview();
});

// ── Settings ───────────────────────────────────────
async function loadSettings() {
  const data = await api('/api/settings');
  document.getElementById('student_name').value = data.student_name || '';
  document.getElementById('student_number').value = data.student_number || '';
}

async function saveSettings() {
  const fd = new FormData();
  fd.append('student_name', document.getElementById('student_name').value);
  fd.append('student_number', document.getElementById('student_number').value);
  await api('/api/settings', { method:'POST', body: fd });
  document.getElementById('settingsSaved').textContent = 'Saved.';
  setTimeout(function() { document.getElementById('settingsSaved').textContent = ''; }, 1200);
}

// ── Quiz ───────────────────────────────────────────
function renderQuizQuestion(data) {
  currentQuiz = data;
  document.getElementById('quizBox').style.display = 'block';
  document.getElementById('quizFeedback').textContent = '';
  clearChallengeArea('quizAskArea');
  clearChallengeArea('quizChallengeArea');
  document.getElementById('quizQ').innerHTML = renderContext(currentQuiz.question);
  lastQuizSubmitted = false;
  quizTotalQuestions = data.total || quizTotalQuestions;

  const counter = 'Question ' + quizQuestionNumber + ' of ' + quizTotalQuestions;
  let qEl = document.getElementById('quizCounter');
  if (!qEl) {
    qEl = document.createElement('div');
    qEl.id = 'quizCounter';
    qEl.className = 'quiz-counter';
    document.getElementById('quizQ').parentNode.insertBefore(qEl, document.getElementById('quizQ'));
  }
  qEl.textContent = counter;

  const prevBtn = document.getElementById('quizPrevBtn');
  if (prevBtn) prevBtn.disabled = quizHistoryIndex <= 0;

  const ctxEl = document.getElementById('quizContext');
  const ctxTextEl = document.getElementById('quizContextText');
  const tb = document.getElementById('quizToggleContext');

  const hasCtx = !!(currentQuiz.context && currentQuiz.context.trim());

  quizHintVisible = false;
  ctxEl.style.display = 'none';
  ctxTextEl.innerHTML = hasCtx ? renderContext(currentQuiz.context) : '';

  if (tb) {
    tb.textContent = 'Show hint';
    tb.disabled = !hasCtx;
    tb.style.display = hasCtx ? 'inline-block' : 'none';
  }

  const box = document.getElementById('quizChoices');
  box.innerHTML = '';
  const indices = getShuffledIndices(currentQuiz);
  indices.forEach(function(choiceIdx, displayIdx) {
    const choiceText = currentQuiz.choices[choiceIdx];
    const id = 'q_' + displayIdx;
    const row = document.createElement('div');
    row.className = 'quiz-choice-row';
    row.innerHTML = '<label class="quiz-choice">' +
      '<input type="radio" name="quizChoice" value="' + choiceIdx + '" id="' + id + '">' +
      '<span class="quiz-choice-text">' + renderContext(choiceText) + '</span>' +
      '</label>';
    box.appendChild(row);
  });
  refreshQuizStats();
}

async function startQuiz() {
  if (!currentUnitId) return alert('Select a unit first.');
  if (requireSubmitBeforeNext && !lastQuizSubmitted && currentQuiz) {
    return alert('Submit your answer before moving to the next question.');
  }

  if (quizHistoryIndex >= 0 && quizHistoryIndex < quizHistory.length - 1) {
    quizHistoryIndex++;
    quizQuestionNumber = quizHistoryIndex + 1;
    renderQuizQuestion(quizHistory[quizHistoryIndex]);
    return;
  }

  const params = new URLSearchParams();
  params.append('unit_id', currentUnitId);
  const recentIds = quizHistory
    .slice(-12) // include last 12 questions to avoid repeats
    .map(q => q.question_id)
    .filter(Boolean);
  if (recentIds.length) {
    params.append('exclude_ids', [...new Set(recentIds)].join(','));
  }
  const data = await api('/api/quiz/random?' + params.toString());
  if (!data || !data.question_id) return alert('No quiz questions seeded for this unit yet.');

  quizHistory.push(data);
  quizHistoryIndex = quizHistory.length - 1;
  quizQuestionNumber = quizHistoryIndex + 1;
  renderQuizQuestion(data);
}

function prevQuiz() {
  if (quizHistoryIndex <= 0) return;
  quizHistoryIndex--;
  quizQuestionNumber = quizHistoryIndex + 1;
  lastQuizSubmitted = true;
  renderQuizQuestion(quizHistory[quizHistoryIndex]);
}

function saveQuizContextPref() {
  // deprecated: kept for backward compatibility with older saved prefs
}

function toggleQuizContext() {
  const tb = document.getElementById('quizToggleContext');
  const ctxEl = document.getElementById('quizContext');
  const ctxTextEl = document.getElementById('quizContextText');
  if (!tb || !ctxEl || !ctxTextEl) return;

  const hasCtx = !!(currentQuiz && currentQuiz.context && currentQuiz.context.trim());
  if (!hasCtx) return;

  quizHintVisible = !quizHintVisible;
  tb.textContent = quizHintVisible ? 'Hide hint' : 'Show hint';

  if (quizHintVisible) {
    ctxEl.style.display = 'block';
    ctxTextEl.innerHTML = renderContext(currentQuiz.context);
  } else {
    ctxEl.style.display = 'none';
  }
}

async function submitQuiz() {
  if (!currentQuiz) return;
  if (lastQuizSubmitted) {
    return startQuiz();
  }
  const chosen = document.querySelector('input[name="quizChoice"]:checked');
  if (!chosen) return alert('Pick an answer first.');

  const submitBtn = document.getElementById('quizSubmitBtn');
  const feedbackEl = document.getElementById('quizFeedback');
  toggleButtonLoading(submitBtn, true, 'Checking...');
  if (feedbackEl) {
    feedbackEl.innerHTML = '<div class="sa-grading">Checking answer...</div>';
  }

  const fd = new FormData();
  fd.append('unit_id', currentUnitId);
  fd.append('question_id', currentQuiz.question_id);
  fd.append('chosen_index', chosen.value);

  try {
    const data = await api('/api/quiz/answer', { method:'POST', body: fd });
    let fb = (data.correct ? 'Correct. ' : 'Not quite. ') + data.explanation;
    if (data.card_created) fb += '\n\u2728 Flashcard created from this mistake.';
    if (feedbackEl) {
      feedbackEl.innerHTML = renderContext(fb);
    }
    if (!data.correct) {
      setChallengeTrigger('quizChallengeArea', {
        kind: 'mcq',
        questionId: currentQuiz.question_id,
        userAnswer: chosen.value,
        unitId: currentUnitId,
      });
    } else {
      setChallengeTrigger('quizChallengeArea', null);
    }
    lastQuizSubmitted = true;
    await refreshQuizStats();
  } catch (e) {
    console.error(e);
    if (feedbackEl) {
      feedbackEl.innerHTML = '<div class="sa-feedback-text" style="color:var(--danger);">Error: ' + e.message + '</div>';
    }
    alert('Failed to submit answer.');
  } finally {
    toggleButtonLoading(submitBtn, false);
  }
}

function clearChallengeArea(containerId) {
  const area = document.getElementById(containerId);
  if (area) area.innerHTML = '';
}

function showQuizAskUI() {
  if (!currentQuiz || !currentUnitId) return;
  showChallengeUI('quizAskArea', {
    kind: 'mcq',
    questionId: currentQuiz.question_id,
    userAnswer: '',
    unitId: currentUnitId,
  });
}

function setChallengeTrigger(containerId, options) {
  const container = document.getElementById(containerId);
  if (!container) return;
  container.innerHTML = '';
  if (!options) return;
  const btn = document.createElement('button');
  btn.type = 'button';
  btn.className = 'btn-ghost challenge-trigger';
  btn.textContent = 'Challenge explanation';
  btn.onclick = () => showChallengeUI(containerId, options);
  container.appendChild(btn);
}

function showChallengeUI(containerId, options) {
  const container = document.getElementById(containerId);
  if (!container) return;
  container.innerHTML = `
    <div class="challenge-box">
      <div class="challenge-header">Ask Gemini for clarification</div>
      <textarea id="${containerId}-input" class="challenge-textarea" rows="3" placeholder="Describe what you don’t understand..."></textarea>
      <div class="challenge-actions">
        <button class="btn-secondary" type="button" onclick="sendChallenge('${containerId}')">Send challenge</button>
      </div>
      <div id="${containerId}-result" class="challenge-result"></div>
    </div>
  `;
  container.dataset.challengeKind = options.kind;
  container.dataset.challengeQuestionId = options.questionId;
  container.dataset.challengeUserAnswer = options.userAnswer || '';
  container.dataset.challengeUnitId = options.unitId || currentUnitId;
}

async function sendChallenge(containerId) {
  const container = document.getElementById(containerId);
  if (!container) return;
  const textarea = document.getElementById(containerId + '-input');
  const resultEl = document.getElementById(containerId + '-result');
  if (!textarea || !resultEl) return;
  const text = textarea.value.trim();
  if (!text) {
    resultEl.innerHTML = '<div class="challenge-error">Tell Gemini what you want clarified.</div>';
    return;
  }
  const fd = new FormData();
  fd.append('unit_id', container.dataset.challengeUnitId || currentUnitId);
  fd.append('question_id', container.dataset.challengeQuestionId);
  fd.append('kind', container.dataset.challengeKind);
  fd.append('user_answer', container.dataset.challengeUserAnswer || '');
  fd.append('user_challenge', text);
  textarea.disabled = true;
  const sendBtn = container.querySelector('.challenge-actions button');
  if (sendBtn) sendBtn.disabled = true;
  resultEl.innerHTML = '<div class="challenge-status">Requesting clarification (Gemini may take a few seconds)...</div>';
  try {
    const data = await api('/api/exam/challenge', { method: 'POST', body: fd });
    if (data && data.ok && data.clarification) {
      resultEl.innerHTML = '<div class="challenge-clarification">' + renderMarkdown(renderContext(data.clarification)) + '</div>';
    } else {
      const err = data && (data.error || data.ai_error) ? (data.error || data.ai_error) : 'Clarification request failed.';
      resultEl.innerHTML = '<div class="challenge-error">Error: ' + err + '</div>';
    }
  } catch (e) {
    resultEl.innerHTML = '<div class="challenge-error">Error: ' + e.message + '</div>';
  } finally {
    textarea.disabled = false;
    if (sendBtn) sendBtn.disabled = false;
  }
}

// ── Exam Mode ───────────────────────────────────────
function showExamSetup() {
  const el = document.getElementById('examContent');
  if (!currentUnitId) {
    el.innerHTML = '<p class="text-muted">Select a unit first.</p>';
    return;
  }
  el.innerHTML = `
    <div class="card">
      <h3>Start a mock exam</h3>
      <p class="text-muted">Questions are drawn from your weakest areas first.</p>
      <div class="field-group">
        <label class="field-label">Total questions</label>
        <select id="examCount">
          <option value="10">10 questions</option>
          <option value="20" selected>20 questions</option>
          <option value="30">30 questions</option>
        </select>
      </div>
      <div class="field-group">
        <label class="field-label">Time limit</label>
        <select id="examMinutes">
          <option value="20">20 minutes</option>
          <option value="30" selected>30 minutes</option>
          <option value="45">45 minutes</option>
        </select>
      </div>
      <button id="examStartBtn" onclick="startExam()" style="margin-top:10px;">Start exam</button>
    </div>
  `;
}

async function startExam() {
  if (!currentUnitId) return alert('Select a unit first.');
  const count = parseInt(document.getElementById('examCount').value, 10);
  const minutes = parseInt(document.getElementById('examMinutes').value, 10);
  const fd = new FormData();
  fd.append('unit_id', currentUnitId);
  fd.append('count', String(count));
  fd.append('minutes', String(minutes));
  const startBtn = document.getElementById('examStartBtn');
  toggleButtonLoading(startBtn, true, 'Starting exam...');
  try {
    const data = await api('/api/exam/start', { method: 'POST', body: fd });
    currentExam = data;
    renderExamQuestion();
  } catch (e) {
    console.error(e);
    if (startBtn) toggleButtonLoading(startBtn, false);
    alert('Failed to start exam.');
  }
}

function renderExamQuestion() {
  if (!currentExam) return;
  const el = document.getElementById('examContent');
  const idx = currentExam.current_index;
  const q = currentExam.questions[idx];
  const total = currentExam.questions.length;
  const remainingSeconds = Math.max(0, Math.floor(currentExam.remaining_seconds));
  const mins = Math.floor(remainingSeconds / 60);
  const secs = remainingSeconds % 60;
  const timeStr = mins + ':' + String(secs).padStart(2, '0');

  let body = `
    <div class="exam-header">
      <div>Question ${idx + 1} of ${total}</div>
      <div class="exam-timer">Time left: ${timeStr}</div>
    </div>
  `;
  body += `<div class="card exam-card"><div class="question-text">${renderContext(q.question)}</div>`;
  if (q.kind === 'mcq') {
    body += '<div id="examChoices">';
    const indices = getShuffledIndices(q);
    indices.forEach((choiceIdx) => {
      const choiceText = q.choices[choiceIdx];
      body += '<div class="quiz-choice-row"><label class="quiz-choice">' +
        '<input type="radio" name="examChoice" value="' + choiceIdx + '">' +
        '<span class="quiz-choice-text">' + renderContext(choiceText) + '</span>' +
        '</label></div>';
    });
    body += '</div>';
  } else if (q.kind === 'sa') {
    body += '<textarea id="examAnswer" rows="5" placeholder="Type your answer here..."></textarea>';
  }
  body += `
    <div class="action-bar" style="margin-top:12px;">
      <button onclick="submitExamAnswer()">Submit & Next</button>
      <button class="btn-secondary" onclick="skipExamQuestion()">Skip</button>
      <button class="btn-ghost" onclick="finishExam()">Finish early</button>
    </div>
  </div>
  `;
  el.innerHTML = body;
}

async function submitExamAnswer() {
  if (!currentExam) return;
  const idx = currentExam.current_index;
  const q = currentExam.questions[idx];
  const fd = new FormData();
  fd.append('exam_id', currentExam.exam_id);
  fd.append('index', String(idx));
  if (q.kind === 'mcq') {
    const chosen = document.querySelector('input[name="examChoice"]:checked');
    if (!chosen) return alert('Pick an answer first.');
    fd.append('answer', chosen.value);
  } else {
    const val = (document.getElementById('examAnswer').value || '').trim();
    fd.append('answer', val);
  }
  try {
    const data = await api('/api/exam/answer', { method: 'POST', body: fd });
    currentExam = data;
    if (currentExam.finished) {
      renderExamResults();
    } else {
      renderExamQuestion();
    }
  } catch (e) {
    console.error(e);
    alert('Failed to submit answer.');
  }
}

async function skipExamQuestion() {
  if (!currentExam) return;
  const idx = currentExam.current_index;
  const fd = new FormData();
  fd.append('exam_id', currentExam.exam_id);
  fd.append('index', String(idx));
  fd.append('answer', '');
  try {
    const data = await api('/api/exam/answer', { method: 'POST', body: fd });
    currentExam = data;
    if (currentExam.finished) {
      renderExamResults();
    } else {
      renderExamQuestion();
    }
  } catch (e) {
    console.error(e);
    alert('Failed to skip question.');
  }
}

async function finishExam() {
  if (!currentExam) return;
  const fd = new FormData();
  fd.append('exam_id', currentExam.exam_id);
  try {
    const data = await api('/api/exam/finish', { method: 'POST', body: fd });
    currentExam = data;
    renderExamResults();
  } catch (e) {
    console.error(e);
    alert('Failed to finish exam.');
  }
}

function renderExamResults() {
  const el = document.getElementById('examContent');
  if (!currentExam || !currentExam.summary) {
    el.innerHTML = '<p class="text-muted">No results available.</p>';
    return;
  }
  const s = currentExam.summary;
  let html = `
    <div class="card">
      <h3>Exam results</h3>
      <p><strong>Score:</strong> ${s.correct}/${s.total} correct (${s.pct.toFixed(1)}%)</p>
      <p><strong>Time used:</strong> ${s.time_used_seconds}s</p>
      <button onclick="showExamSetup()" class="btn-secondary">Start another exam</button>
      <button onclick="showPage('mistakes')" class="btn-ghost" style="margin-left:8px;">Review mistakes</button>
    </div>
  `;
  html += '<div class="card" style="margin-top:12px;"><h4>Question review</h4>';
  s.questions.forEach((q, i) => {
    const reviewBodyId = 'examReviewBody-' + i;
    const challengeAreaId = 'examChallengeArea-' + i;
    html += '<div class="exam-review-item">';
    html += '<div class="exam-review-header"><strong>Q' + (i + 1) + '.</strong> ' + renderContext(q.question) +
      ' <span class="' + (q.correct ? 'badge-correct' : 'badge-incorrect') + '">' + (q.correct ? 'Correct' : 'Wrong') + '</span></div>';
    if (q.kind === 'mcq') {
      html += '<div id="' + reviewBodyId + '" class="exam-review-body"><div><strong>Your answer:</strong> ' + renderContext(q.your_answer || '(blank)') + '</div>' +
              '<div><strong>Correct answer:</strong> ' + renderContext(q.correct_answer) + '</div></div>';
    } else {
      html += '<div id="' + reviewBodyId + '" class="exam-review-body"><div><strong>Your answer:</strong><br>' + renderContext(q.your_answer || '(blank)') + '</div>' +
              '<div style="margin-top:4px;"><strong>Model answer:</strong><br>' + renderContext(q.correct_answer) + '</div>' +
              (q.ai_feedback ? '<div style="margin-top:4px;color:var(--text-muted);"><strong>Feedback:</strong> ' + renderContext(q.ai_feedback) + '</div>' : '') + '</div>';
    }
    html += '<div id="' + challengeAreaId + '" class="challenge-area"></div>';
    if (!q.correct) {
      html += '<div class="exam-review-actions"><button class="btn-ghost" onclick="launchExamChallenge(\'' + challengeAreaId + '\',' + i + ')">Challenge explanation</button></div>';
    }
    html += '</div>';
  });
  html += '</div>';
  el.innerHTML = html;
}

function launchExamChallenge(containerId, questionIndex) {
  if (!currentExam || !currentExam.summary) return;
  const q = currentExam.summary.questions[questionIndex];
  if (!q) return;
  const reviewBody = document.getElementById('examReviewBody-' + questionIndex);
  if (reviewBody) reviewBody.style.display = 'none';
  showChallengeUI(containerId, {
    kind: q.kind,
    questionId: q.question_id,
    userAnswer: q.your_answer,
    unitId: currentExam.summary.unit_id || currentUnitId,
  });
}

// ── Explain-It-Back ────────────────────────────────
async function loadExplainTopic() {
  const el = document.getElementById('explainContent');
  if (!currentUnitId) {
    el.innerHTML = '<p class="text-muted">Select a unit first.</p>';
    return;
  }
  try {
    const data = await api('/api/explain/random?unit_id=' + currentUnitId);
    if (!data || !data.topic_id) {
      el.innerHTML = '<p class="text-muted">No explain-it-back topics seeded for this unit yet.</p>';
      return;
    }
    currentExplainTopic = data;
    renderExplainTopic();
  } catch (e) {
    console.error(e);
    el.innerHTML = '<p class="text-muted">Failed to load topic.</p>';
  }
}

function renderExplainTopic() {
  const el = document.getElementById('explainContent');
  if (!currentExplainTopic) {
    el.innerHTML = '<p class="text-muted">No topic loaded.</p>';
    return;
  }
  const t = currentExplainTopic;
  el.innerHTML = `
    <div class="card">
      <h3>Explain this concept</h3>
      <p class="question-text">${renderContext(t.topic_prompt)}</p>
      <textarea id="explainAnswer" rows="6" placeholder="Explain it in your own words..."></textarea>
      <div class="action-bar" style="margin-top:12px;">
        <button onclick="submitExplain()">Check my explanation</button>
        <button class="btn-secondary" onclick="loadExplainTopic()">New topic</button>
      </div>
      <div id="explainFeedback" class="feedback-area"></div>
    </div>
  `;
}

async function submitExplain() {
  if (!currentExplainTopic) return;
  const txt = (document.getElementById('explainAnswer').value || '').trim();
  if (!txt) return alert('Write your explanation first.');
  const fd = new FormData();
  fd.append('unit_id', currentUnitId);
  fd.append('topic_id', currentExplainTopic.topic_id);
  fd.append('student_text', txt);
  const fb = document.getElementById('explainFeedback');
  fb.textContent = 'Checking...';
  try {
    const data = await api('/api/explain/check', { method: 'POST', body: fd });
    let stars = '';
    for (let i = 0; i < 5; i++) {
      stars += i < data.ai_score ? '★' : '☆';
    }
    let html = '<div class="sa-verdict"><span class="sa-verdict-partial">' + stars + '</span></div>';
    html += '<div class="sa-feedback-text">' + renderContext(data.ai_feedback || '') + '</div>';
    html += '<div class="sa-model-answer" style="margin-top:8px;"><strong>Model explanation:</strong><br>' + renderContext(data.model_explanation) + '</div>';
    fb.innerHTML = html;
  } catch (e) {
    console.error(e);
    fb.textContent = 'AI check failed. Try again later.';
  }
}
// ── Mistake review ─────────────────────────────────
async function loadMistakes() {
  if (!currentUnitId) {
    document.getElementById('mistakeList').textContent = 'Select a unit first.';
    return;
  }
  try {
    const data = await api('/api/mistakes?unit_id=' + currentUnitId);
    mistakesData = data;
    renderMistakes();
  } catch (e) {
    console.error(e);
    document.getElementById('mistakeList').textContent = 'Failed to load mistakes.';
  }
}

function setMistakeFilter(tag) {
  mistakeFilterTag = tag || '';
  try {
    localStorage.setItem('mistakeFilterTag', mistakeFilterTag);
  } catch (e) {
    console.error('Failed to persist mistake filter:', e);
  }
  renderMistakes();
}

function renderMistakes() {
  const list = document.getElementById('mistakeList');
  const tagBar = document.getElementById('mistakeTags');
  if (!mistakesData) {
    list.textContent = 'No mistakes loaded yet.';
    tagBar.innerHTML = '';
    return;
  }

  const allItems = []
    .concat((mistakesData.quiz || []).map(m => ({ type: 'quiz', ...m })))
    .concat((mistakesData.short_answer || []).map(m => ({ type: 'sa', ...m })));

  if (!allItems.length) {
    list.textContent = 'Nice work — no recorded mistakes yet.';
    tagBar.innerHTML = '';
    return;
  }

  // Build tag filter list
  const tagSet = new Set();
  allItems.forEach(m => {
    (m.tags || '').split(',').map(t => t.trim()).filter(Boolean).forEach(t => tagSet.add(t));
  });
  const tags = Array.from(tagSet).sort();
  let tagHtml = '<span class="tag-chip ' + (!mistakeFilterTag ? 'tag-chip-active' : '') + '" onclick="setMistakeFilter(\'\')">All</span>';
  tags.forEach(t => {
    const active = mistakeFilterTag === t;
    tagHtml += '<span class="tag-chip ' + (active ? 'tag-chip-active' : '') + '" onclick="setMistakeFilter(\'' + t.replace(/'/g, "\\'") + '\')">' + t + '</span>';
  });
  tagBar.innerHTML = tagHtml;

  const filtered = mistakeFilterTag
    ? allItems.filter(m => (m.tags || '').split(',').map(t => t.trim()).includes(mistakeFilterTag))
    : allItems;

  if (!filtered.length) {
    list.textContent = 'No mistakes for this tag yet.';
    return;
  }

  list.innerHTML = '';
  filtered.sort((a, b) => (a.last_ts < b.last_ts ? 1 : -1));

  filtered.forEach(m => {
    const card = document.createElement('div');
    card.className = 'card mistake-card';
    const masteredBadge = m.mastered ? '<span class="mistake-badge mastered">Mastered</span>' : '<span class="mistake-badge pending">Still tricky</span>';
    const tags = (m.tags || '').split(',').map(t => t.trim()).filter(Boolean).join(', ');
    const sourceLabel = m.source === 'mcq_quiz' ? 'MCQ Quiz' : m.source === 'short_answer' ? 'Short Answer' : m.source === 'exam' ? 'Exam Mode' : m.source;
    let body = '<div class="mistake-header"><div class="mistake-meta">' + masteredBadge +
      '<span class="mistake-count">Wrong ' + m.wrong_count + '×</span>' +
      '<span class="mistake-source">' + sourceLabel + '</span>' +
      (tags ? '<span class="mistake-tags">' + tags + '</span>' : '') +
      '</div></div>';
    body += '<div class="mistake-question">' + renderContext(m.question) + '</div>';
    if (m.type === 'quiz') {
      const last = m.last_chosen_index;
      const correct = m.correct_index;
      const lastText = m.choices[last] != null ? m.choices[last] : '(unknown choice)';
      const correctText = m.choices[correct] != null ? m.choices[correct] : '(unknown choice)';
      body += '<div class="mistake-detail"><div><strong>Your last answer:</strong> ' + renderContext(lastText) + '</div>' +
              '<div><strong>Correct answer:</strong> ' + renderContext(correctText) + '</div></div>';
      body += '<div class="mistake-actions"><button onclick=\'reattemptQuizMistake(' +
              JSON.stringify({ question_id: m.question_id, question: m.question, choices: m.choices, context: m.context || '' }).replace(/'/g, "\\'") +
              ")'>Re-attempt as quiz</button></div>";
    } else {
      body += '<div class="mistake-detail"><div><strong>Your last answer:</strong><br>' + renderContext(m.last_student_answer || '(blank)') + '</div>' +
              '<div style="margin-top:6px;"><strong>Model answer:</strong><br>' + renderContext(m.model_answer) + '</div></div>';
      body += '<div class="mistake-actions"><button onclick=\'reattemptSAMistake(' +
              JSON.stringify({ question_id: m.question_id, question: m.question, context: m.context || '' }).replace(/'/g, "\\'") +
              ")'>Re-attempt as short answer</button></div>";
    }
    card.innerHTML = body;
    list.appendChild(card);
  });
}

function reattemptQuizMistake(payload) {
  if (!payload) return;
  showPage('quiz');
  currentQuiz = {
    question_id: payload.question_id,
    question: payload.question,
    choices: payload.choices,
    context: payload.context || '',
    total: 0
  };
  quizHistory.push(currentQuiz);
  quizHistoryIndex = quizHistory.length - 1;
  quizQuestionNumber = quizHistoryIndex + 1;
  lastQuizSubmitted = false;
  renderQuizQuestion(currentQuiz);
}

function reattemptSAMistake(payload) {
  if (!payload) return;
  showPage('short-answer');
  currentSA = {
    question_id: payload.question_id,
    question: payload.question,
    context: payload.context || '',
    total: 0
  };
  saHistory.push(currentSA);
  saHistoryIndex = saHistory.length - 1;
  saQuestionNumber = saHistoryIndex + 1;
  lastSASubmitted = false;
  renderSAQuestion(currentSA);
}

// ── Glossary ───────────────────────────────────────
const GLOSSARY = {
  "SSH": "Secure Shell \u2014 a protocol for encrypted remote administration over a network.",
  "Telnet": "An older remote terminal protocol that sends all data (including passwords) in plaintext.",
  "HTTPS": "HTTP Secure \u2014 HTTP encrypted with TLS, protecting data between browser and server.",
  "HTTP": "HyperText Transfer Protocol \u2014 the foundation of web communication (unencrypted by default).",
  "DNS": "Domain Name System \u2014 translates domain names (e.g. google.com) into IP addresses.",
  "DHCP": "Dynamic Host Configuration Protocol \u2014 automatically assigns IP addresses to devices on a network.",
  "SNMP": "Simple Network Management Protocol \u2014 used to monitor and manage network devices like routers and switches.",
  "FTP": "File Transfer Protocol \u2014 transfers files between systems; sends credentials in plaintext unless using SFTP/FTPS.",
  "TCP": "Transmission Control Protocol \u2014 a reliable, connection-oriented transport protocol that ensures data arrives in order.",
  "UDP": "User Datagram Protocol \u2014 a fast, connectionless transport protocol with no guaranteed delivery.",
  "IP": "Internet Protocol \u2014 the addressing protocol that routes packets across networks using IP addresses.",
  "TLS": "Transport Layer Security \u2014 a cryptographic protocol that encrypts communication (successor to SSL).",
  "SSL": "Secure Sockets Layer \u2014 an older encryption protocol replaced by TLS; the term is still commonly used.",
  "ICMP": "Internet Control Message Protocol \u2014 used for network diagnostics (e.g. ping, traceroute).",
  "ACL": "Access Control List \u2014 a set of rules on a router or firewall that permit or deny traffic based on IP, port, or protocol.",
  "DMZ": "Demilitarised Zone \u2014 an isolated network segment for public-facing servers, protecting the internal LAN.",
  "VPN": "Virtual Private Network \u2014 creates an encrypted tunnel over a public network for secure remote access.",
  "IDS": "Intrusion Detection System \u2014 monitors network traffic for suspicious activity and generates alerts.",
  "IPS": "Intrusion Prevention System \u2014 like an IDS but can also actively block or prevent detected threats.",
  "NIDS": "Network-based IDS \u2014 monitors traffic at a network tap or span port to detect threats across a segment.",
  "HIDS": "Host-based IDS \u2014 monitors activity on a single host/server for signs of compromise.",
  "MFA": "Multi-Factor Authentication \u2014 requires two or more different factor types: knowledge, possession, or biometric.",
  "firewall": "A system that filters network traffic according to security rules, controlling what enters and leaves a network.",
  "proxy": "A server that acts as an intermediary between clients and destination servers, often for filtering or caching.",
  "NAT": "Network Address Translation \u2014 maps private internal IP addresses to a public IP for internet access.",
  "VLAN": "Virtual LAN \u2014 a logical network segment that separates traffic within the same physical switch.",
  "hash": "A fixed-size digest computed from data using a one-way function; used to verify data integrity.",
  "encryption": "The process of converting readable data into unreadable ciphertext using a key; provides confidentiality.",
  "checksum": "A computed value used to verify that data has not been altered or corrupted during transfer or storage.",
  "RAID": "Redundant Array of Independent Disks \u2014 provides hardware redundancy but is NOT a substitute for backups.",
  "backup": "A copy of data stored separately so it can be restored after loss, corruption, or disaster.",
  "RPO": "Recovery Point Objective \u2014 the maximum acceptable amount of data loss, measured in time.",
  "RTO": "Recovery Time Objective \u2014 the maximum acceptable downtime before services must be restored.",
  "DRP": "Disaster Recovery Plan \u2014 documented procedures for restoring IT systems after a major disruption.",
  "GFS": "Grandfather-Father-Son \u2014 a backup rotation scheme using daily, weekly, and monthly backup cycles.",
  "UPS": "Uninterruptible Power Supply \u2014 provides battery power during outages for safe shutdown or generator bridging.",
  "MITM": "Man-in-the-Middle \u2014 an attack where an attacker secretly intercepts and may alter communication between two parties.",
  "eavesdropping": "A passive attack that captures network traffic without altering it (e.g. packet sniffing).",
  "phishing": "A social engineering attack using fake emails or websites to trick users into revealing credentials or data.",
  "ransomware": "Malware that encrypts a victim's files and demands payment for the decryption key.",
  "brute-force": "An attack that systematically tries all possible passwords or keys until the correct one is found.",
  "DNS poisoning": "Injecting false DNS records into a resolver's cache, redirecting users to malicious sites.",
  "Privacy Act": "Australian Privacy Act 1988 \u2014 governs how organisations collect, use, store, and disclose personal information.",
  "APPs": "Australian Privacy Principles \u2014 13 principles under the Privacy Act governing personal information handling.",
  "ISO 27001": "An international standard for information security management systems (ISMS).",
  "penetration test": "A simulated attack on a system to find exploitable vulnerabilities before real attackers do.",
  "vulnerability scan": "An automated tool that checks systems for known weaknesses like missing patches or misconfigurations.",
  "least privilege": "The principle of giving users/services only the minimum access needed to perform their tasks.",
  "defence in depth": "A security strategy using multiple overlapping layers of controls so no single failure is catastrophic.",
  "risk assessment": "The process of identifying threats, assessing likelihood and impact, and prioritising risks."
};

function escapeRegExp(string) {
  return string.replace(/[.*+?^${}()|[\\]\\]/g, '\\$&');
}

function renderContext(text) {
  let out = text;
  Object.keys(GLOSSARY).forEach(function(term) {
    const re = new RegExp('\\b' + escapeRegExp(term) + '\\b', 'gi');
    out = out.replace(re, function(match) { return '<a href="#" class="glossary-term" data-term="' + term + '">' + match + '</a>'; });
  });
  out = out.replace(/\n/g, '<br>');
  return out;
}

function renderMarkdown(text) {
  if (!text) return '';
  return text
    .replace(/\*\*(.+?)\*\*/g, '<strong>$1</strong>')
    .replace(/\*(?!\*)(.+?)\*(?!\*)/g, '<em>$1</em>');
}

document.addEventListener('click', function(ev) {
  const a = ev.target.closest && ev.target.closest('.glossary-term');
  if (a) {
    ev.preventDefault();
    const term = a.getAttribute('data-term');
    openGlossary(term);
  }
});

function openGlossary(term) {
  const modal = document.getElementById('glossaryModal');
  const title = document.getElementById('glossaryTitle');
  const body = document.getElementById('glossaryBody');
  if (!modal) return;
  title.textContent = term;
  body.textContent = GLOSSARY[term] || 'Definition not found.';
  modal.style.display = 'block';
}

function closeGlossary() {
  const modal = document.getElementById('glossaryModal');
  if (modal) modal.style.display = 'none';
}

// ── Short Answer ───────────────────────────────────
function renderSAQuestion(data) {
  currentSA = data;
  document.getElementById('saBox').style.display = 'block';
  document.getElementById('saFeedback').innerHTML = '';
  clearChallengeArea('saChallengeArea');
  document.getElementById('saModelAnswer').style.display = 'none';
  document.getElementById('saQ').innerHTML = renderContext(currentSA.question);
  document.getElementById('saAnswer').value = '';
  document.getElementById('saAnswer').disabled = false;
  document.getElementById('saSubmitBtn').disabled = false;
  document.getElementById('saTeachBtn').disabled = false;
  lastSASubmitted = false;
  saTotalQuestions = data.total || saTotalQuestions;

  const counter = 'Question ' + saQuestionNumber + ' of ' + saTotalQuestions;
  let cEl = document.getElementById('saCounter');
  if (!cEl) {
    cEl = document.createElement('div');
    cEl.id = 'saCounter';
    cEl.className = 'quiz-counter';
    document.getElementById('saQ').parentNode.insertBefore(cEl, document.getElementById('saQ'));
  }
  cEl.textContent = counter;

  const prevBtn = document.getElementById('saPrevBtn');
  if (prevBtn) prevBtn.disabled = saHistoryIndex <= 0;

  const ctxEl = document.getElementById('saContext');
  const ctxTextEl = document.getElementById('saContextText');
  const tb = document.getElementById('saToggleContext');
  const hasCtx = !!(currentSA.context && currentSA.context.trim());

  saHintVisible = false;
  ctxEl.style.display = 'none';
  ctxTextEl.innerHTML = hasCtx ? renderContext(currentSA.context) : '';

  if (tb) {
    tb.textContent = 'Show hint';
    tb.disabled = !hasCtx;
    tb.style.display = hasCtx ? 'inline-block' : 'none';
  }
  refreshSAStats();
}

async function startShortAnswer() {
  if (!currentUnitId) return alert('Select a unit first.');

  if (saHistoryIndex >= 0 && saHistoryIndex < saHistory.length - 1) {
    saHistoryIndex++;
    saQuestionNumber = saHistoryIndex + 1;
    renderSAQuestion(saHistory[saHistoryIndex]);
    return;
  }

  const data = await api('/api/short-answer/random?unit_id=' + currentUnitId);
  if (!data || !data.question_id) return alert('No short answer questions seeded for this unit yet.');

  saHistory.push(data);
  saHistoryIndex = saHistory.length - 1;
  saQuestionNumber = saHistoryIndex + 1;
  renderSAQuestion(data);
}

function prevShortAnswer() {
  if (saHistoryIndex <= 0) return;
  saHistoryIndex--;
  saQuestionNumber = saHistoryIndex + 1;
  lastSASubmitted = true;
  renderSAQuestion(saHistory[saHistoryIndex]);
}

function toggleSAContext() {
  const tb = document.getElementById('saToggleContext');
  const ctxEl = document.getElementById('saContext');
  const ctxTextEl = document.getElementById('saContextText');
  if (!tb || !ctxEl || !ctxTextEl) return;

  const hasCtx = !!(currentSA && currentSA.context && currentSA.context.trim());
  if (!hasCtx) return;

  saHintVisible = !saHintVisible;
  tb.textContent = saHintVisible ? 'Hide hint' : 'Show hint';

  if (saHintVisible) {
    ctxEl.style.display = 'block';
    ctxTextEl.innerHTML = renderContext(currentSA.context);
  } else {
    ctxEl.style.display = 'none';
  }
}

async function teachMeSA() {
  if (!currentSA || lastSASubmitted) return;
  if (!currentUnitId) return alert('Select a unit first.');

  const feedbackEl = document.getElementById('saFeedback');
  feedbackEl.innerHTML = '<div class="sa-grading">Learning about this topic...</div>';
  document.getElementById('saSubmitBtn').disabled = true;
  document.getElementById('saTeachBtn').disabled = true;
  document.getElementById('saAnswer').disabled = true;

  const fd = new FormData();
  fd.append('unit_id', currentUnitId);
  fd.append('question_id', currentSA.question_id);

  try {
    const data = await api('/api/short-answer/teach', { method: 'POST', body: fd });

    let html = '';
    if (data.teaching) {
      html += '<div class="sa-teach-box"><div class="sa-teach-label">Here\'s what you need to know:</div>' + renderContext(data.teaching) + '</div>';
    }
    if (data.card_created) {
      html += '<div class="card-created-note">Flashcard created for review.</div>';
    }
    feedbackEl.innerHTML = html;

    document.getElementById('saModelAnswer').style.display = 'block';
    document.getElementById('saModelAnswerText').textContent = data.model_answer;
    lastSASubmitted = true;
  } catch (e) {
    feedbackEl.innerHTML = '<div class="sa-feedback-text" style="color:var(--danger);">Error: ' + e.message + '</div>';
    document.getElementById('saSubmitBtn').disabled = false;
    document.getElementById('saTeachBtn').disabled = false;
    document.getElementById('saAnswer').disabled = false;
  }

  await refreshSAStats();
}

async function submitShortAnswer() {
  if (!currentSA) return;
  if (lastSASubmitted) {
    return startShortAnswer();
  }
  const answer = document.getElementById('saAnswer').value.trim();
  if (!answer) return alert('Type an answer first.');

  const feedbackEl = document.getElementById('saFeedback');
  feedbackEl.innerHTML = '<div class="sa-grading">Grading with Gemini (if configured)...</div>';
  document.getElementById('saSubmitBtn').disabled = true;

  const fd = new FormData();
  fd.append('unit_id', currentUnitId);
  fd.append('question_id', currentSA.question_id);
  fd.append('student_answer', answer);

  try {
    const data = await api('/api/short-answer/check', { method: 'POST', body: fd });

    if (data.mode === 'self_grade') {
      document.getElementById('saModelAnswer').style.display = 'block';
      document.getElementById('saModelAnswerText').textContent = data.model_answer;
      const reason = data.ai_error
        ? data.ai_error + ' \u2014 Compare your answer to the model answer and grade yourself:'
        : 'No API key set \u2014 compare your answer to the model answer and grade yourself:';
      feedbackEl.innerHTML =
        '<div class="sa-feedback-text">' + reason + '</div>' +
        '<div class="sa-self-grade row">' +
        '<button onclick="selfGradeSA(2)">Correct</button>' +
        '<button onclick="selfGradeSA(1)">Partial</button>' +
        '<button onclick="selfGradeSA(0)">Incorrect</button>' +
        '</div>';
      setChallengeTrigger('saChallengeArea', {
        kind: 'sa',
        questionId: currentSA.question_id,
        userAnswer: answer,
        unitId: currentUnitId,
      });
    } else {
      const verdictClass = data.score === 2 ? 'sa-verdict-correct' : data.score === 1 ? 'sa-verdict-partial' : 'sa-verdict-incorrect';
      const verdictText = data.score === 2 ? 'Correct' : data.score === 1 ? 'Partial' : 'Incorrect';
      let cardNote = data.card_created ? '<div class="card-created-note">Flashcard created from this mistake.</div>' : '';
      feedbackEl.innerHTML =
        '<span class="sa-verdict ' + verdictClass + '">' + verdictText + '</span>' +
        '<div class="sa-feedback-text">' + (data.ai_feedback || '') + '</div>' +
        cardNote;
      document.getElementById('saModelAnswer').style.display = 'block';
      document.getElementById('saModelAnswerText').textContent = data.model_answer;
      lastSASubmitted = true;
      if (data.score < 2) {
        setChallengeTrigger('saChallengeArea', {
          kind: 'sa',
          questionId: currentSA.question_id,
          userAnswer: answer,
          unitId: currentUnitId,
        });
      } else {
        setChallengeTrigger('saChallengeArea', null);
      }
    }
  } catch (e) {
    feedbackEl.innerHTML = '<div class="sa-feedback-text" style="color:var(--danger);">Error: ' + e.message + '</div>';
    document.getElementById('saSubmitBtn').disabled = false;
  }

  document.getElementById('saAnswer').disabled = true;
  await refreshSAStats();
}

async function selfGradeSA(score) {
  if (!currentSA) return;
  const fd = new FormData();
  fd.append('unit_id', currentUnitId);
  fd.append('question_id', currentSA.question_id);
  fd.append('score', score);
  fd.append('student_answer', document.getElementById('saAnswer').value.trim());

  const sgData = await api('/api/short-answer/self-grade', { method: 'POST', body: fd });

  const verdictClass = score === 2 ? 'sa-verdict-correct' : score === 1 ? 'sa-verdict-partial' : 'sa-verdict-incorrect';
  const verdictText = score === 2 ? 'Correct' : score === 1 ? 'Partial' : 'Incorrect';
  let cardNote = (sgData && sgData.card_created) ? '<div class="card-created-note">Flashcard created from this mistake.</div>' : '';
  document.getElementById('saFeedback').innerHTML =
    '<span class="sa-verdict ' + verdictClass + '">' + verdictText + '</span>' +
    '<div class="sa-feedback-text">Self-graded.</div>' +
    cardNote;
  lastSASubmitted = true;
  if (score < 2) {
    setChallengeTrigger('saChallengeArea', {
      kind: 'sa',
      questionId: currentSA.question_id,
      userAnswer: document.getElementById('saAnswer').value.trim(),
      unitId: currentUnitId,
    });
  } else {
    setChallengeTrigger('saChallengeArea', null);
  }
  await refreshSAStats();
}

async function refreshSAStats() {
  if (!currentUnitId) return;
  const el = document.getElementById('saStats');
  if (!el) return;

  const s = await api('/api/short-answer/stats?unit_id=' + currentUnitId + '&window=20');
  if (!s.total) { el.innerHTML = ''; return; }

  let html = '<div class="quiz-stats-header">' +
    '<div class="quiz-stats-score">' +
    '<span class="quiz-stats-big">' + s.pct + '%</span>' +
    '<span class="quiz-stats-detail">' + s.correct + '/' + s.total + ' correct (last ' + s.window + ')</span>' +
    '</div>' +
    '<div class="quiz-stats-bar-track">' +
    '<div class="quiz-stats-bar-fill" style="width:' + Math.min(s.pct, 100) + '%"></div>' +
    '</div></div>';

  if (s.by_tag && s.by_tag.length) {
    const weakTags = s.by_tag.filter(function(t) { return t.pct < 100; }).slice(0, 5);
    if (weakTags.length) {
      html += '<div class="quiz-stats-tags">';
      weakTags.forEach(function(t) {
        html += '<span class="quiz-stats-tag">' + t.tag + ' <strong>' + t.pct + '%</strong></span>';
      });
      html += '</div>';
    }
  }
  el.innerHTML = html;
}

// ── Progress ───────────────────────────────────────
async function loadProgress() {
  if (!currentUnitId) return;
  const el = document.getElementById('progressDash');
  if (!el) return;

  const p = await api('/api/progress?unit_id=' + currentUnitId);
  if (!p.days || p.days.length === 0) {
    el.innerHTML = '<div class="text-muted">Answer some questions to see your progress here.</div>';
    return;
  }

  let html = '<div class="progress-header">';
  html += '<div class="progress-stat"><span class="progress-stat-big">' + p.streak + '</span><span class="progress-stat-label">Day streak</span></div>';
  html += '<div class="progress-stat"><span class="progress-stat-big">' + p.total_answered + '</span><span class="progress-stat-label">Total answered</span></div>';
  html += '<div class="progress-stat"><span class="progress-stat-big">' + p.overall_pct + '%</span><span class="progress-stat-label">Overall</span></div>';
  html += '</div>';

  var maxQ = Math.max(1, Math.max.apply(null, p.days.map(function(d) { return d.total; })));
  html += '<div class="progress-chart-wrap"><div class="progress-chart">';
  p.days.forEach(function(d) {
    var h = Math.max(2, (d.total / maxQ) * 100);
    var pct = d.total > 0 ? Math.round((d.correct / d.total) * 100) : 0;
    var opacity = d.total > 0 ? 0.4 + (pct / 100) * 0.6 : 0.1;
    var label = d.date.slice(5);
    html += '<div class="progress-bar" style="height:100%;" title="' + d.date + ': ' + d.correct + '/' + d.total + ' (' + pct + '%)">' +
      '<div class="progress-bar-fill" style="height:' + h + '%; opacity:' + opacity + ';"></div>' +
      '<span class="progress-bar-label">' + label + '</span>' +
      '</div>';
  });
  html += '</div></div>';

  el.innerHTML = html;
}

// ── Gemini Key ─────────────────────────────────────
async function saveGeminiKey() {
  const key = document.getElementById('settings_gemini_key').value.trim();
  const fd = new FormData();
  fd.append('gemini_api_key', key);
  await api('/api/settings/gemini-key', { method: 'POST', body: fd });
  document.getElementById('geminiKeySaved').textContent = 'Saved.';
  setTimeout(function() { document.getElementById('geminiKeySaved').textContent = ''; }, 1200);
}

async function loadGeminiKeyStatus() {
  const data = await api('/api/settings/gemini-key-status');
  if (data.has_key) {
    document.getElementById('settings_gemini_key').placeholder = 'Key saved (hidden). Paste a new one to replace.';
  }
}

// ── AE2 Case Study ─────────────────────────────────
async function loadAE2Items() {
  const sel = document.getElementById('ae2ItemSelect');
  sel.innerHTML = '';
  if (!currentUnitId) return;

  const data = await api('/api/ae2/items?unit_id=' + currentUnitId);
  data.items.forEach(function(it) {
    const opt = document.createElement('option');
    opt.value = it.code;
    opt.textContent = it.code + ' \u2014 ' + it.title;
    sel.appendChild(opt);
  });

  if (data.items.length > 0) {
    sel.value = data.items[0].code;
    await loadAE2Item();
  }
}

async function loadAE2Item() {
  const code = document.getElementById('ae2ItemSelect').value;
  if (!currentUnitId || !code) return;

  const data = await api('/api/ae2/item?unit_id=' + currentUnitId + '&item_code=' + encodeURIComponent(code));
  document.getElementById('ae2Content').value = data.content_md || data.template_md || '';
  document.getElementById('ae2Status').textContent = data.word_guidance ? 'Word guidance: ~' + data.word_guidance + ' words.' : '';
}

async function saveAE2() {
  const code = document.getElementById('ae2ItemSelect').value;
  const content = document.getElementById('ae2Content').value;

  const fd = new FormData();
  fd.append('unit_id', currentUnitId);
  fd.append('item_code', code);
  fd.append('content_md', content);

  await api('/api/ae2/response/save', { method:'POST', body: fd });
  document.getElementById('ae2Status').textContent = 'Saved.';
  setTimeout(function() { loadAE2Item(); }, 200);
}

function initAE2Autosave() {
  const textarea = document.getElementById('ae2Content');
  if (!textarea) return;

  textarea.addEventListener('input', function() {
    const status = document.getElementById('ae2Status');
    if (status) status.textContent = 'Saving...';

    if (ae2AutosaveTimer) {
      clearTimeout(ae2AutosaveTimer);
    }
    ae2AutosaveTimer = setTimeout(function() {
      saveAE2().catch(function() {
        if (status) status.textContent = 'Autosave failed.';
      });
    }, 1000);
  });
}

function initSettingsUI() {
  const checkbox = document.getElementById('settings_require_submit');
  if (!checkbox) return;

  const raw = localStorage.getItem('requireSubmitBeforeNext');
  requireSubmitBeforeNext = raw === 'true';
  checkbox.checked = requireSubmitBeforeNext;

  checkbox.addEventListener('change', function() {
    requireSubmitBeforeNext = checkbox.checked;
    localStorage.setItem('requireSubmitBeforeNext', requireSubmitBeforeNext ? 'true' : 'false');
  });
}

async function restoreAppState() {
  const savedMode = localStorage.getItem('flashcardMode');
  if (savedMode) {
    currentMode = savedMode;
    const cardKindEl = document.getElementById('card_event_kind');
    if (cardKindEl) cardKindEl.value = savedMode;
  }

  const savedUnitId = localStorage.getItem('currentUnitId');
  if (savedUnitId && availableUnits.length) {
    const match = availableUnits.find(u => String(u.id) === savedUnitId);
    if (match) {
      await selectUnit(match);
    }
  }

  const savedPage = localStorage.getItem('selectedPage') || 'dashboard';
  showPage(savedPage);
}

async function exportAE2() {
  const code = document.getElementById('ae2ItemSelect').value;
  const url = '/api/ae2/export_docx?unit_id=' + currentUnitId + '&item_code=' + encodeURIComponent(code);
  const res = await fetch(url);
  if (!res.ok) return alert('Export failed.');
  const blob = await res.blob();
  const cd = res.headers.get('content-disposition') || '';
  let filename = 'AE2.docx';
  const m = cd.match(/filename="([^"]+)"/);
  if (m) filename = m[1];

  const a = document.createElement('a');
  a.href = URL.createObjectURL(blob);
  a.download = filename;
  document.body.appendChild(a);
  a.click();
  a.remove();
}

// ── Init ───────────────────────────────────────────
unitsLoadedPromise = loadUnits();
document.addEventListener('DOMContentLoaded', function() {
  unitsLoadedPromise
    .then(() => restoreAppState())
    .catch(function(e) { console.error('Failed to restore UI state:', e); });
  initAE2Autosave();
  initSettingsUI();
  loadGeminiKeyStatus();
});
"""


@app.on_event("startup")
def startup():
    migrate()
    seed_all()


@app.get("/", response_class=HTMLResponse)
def index():
    return INDEX_HTML


@app.get("/styles.css")
def styles():
    return Response(content=STYLES_CSS, media_type="text/css")


@app.get("/app.js")
def appjs():
    return Response(content=APP_JS, media_type="application/javascript")


@app.get("/api/settings")
def get_settings():
    con = connect()
    rows = con.execute("SELECT key, value FROM settings;").fetchall()
    con.close()
    kv = {r["key"]: r["value"] for r in rows}
    return {"student_name": kv.get("student_name", ""), "student_number": kv.get("student_number", "")}


@app.post("/api/settings")
def set_settings(student_name: str = Form(""), student_number: str = Form("")):
    con = connect()
    con.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?);", ("student_name", student_name.strip()))
    con.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?);", ("student_number", student_number.strip()))
    con.commit()
    con.close()
    return {"ok": True}


@app.get("/api/units")
def get_units():
    con = connect()
    rows = con.execute(
        "SELECT id, code, title, pinned FROM units ORDER BY pinned DESC, code ASC;"
    ).fetchall()
    con.close()
    return {"units": [dict(r) for r in rows]}


@app.post("/api/cards")
def create_card(
    unit_id: int = Form(...),
    event_kind: str = Form(...),
    prompt: str = Form(...),
    answer: str = Form(...),
    tags: str = Form("")
):
    con = connect()
    cur = con.cursor()
    cur.execute(
        "INSERT INTO cards (unit_id, event_kind, prompt, answer, tags) VALUES (?, ?, ?, ?, ?);",
        (unit_id, event_kind.strip(), prompt.strip(), answer.strip(), tags.strip())
    )
    card_id = cur.lastrowid
    cur.execute(
        "INSERT OR REPLACE INTO reviews (card_id, due_date, interval_days, ease) VALUES (?, ?, ?, ?);",
        (card_id, date.today().isoformat(), 1, 2.5)
    )
    con.commit()
    con.close()
    return {"ok": True, "card_id": card_id}


@app.get("/api/review/today")
def review_today(unit_id: int, event_kind: str):
    con = connect()
    rows = con.execute("""
        SELECT c.id, c.prompt, c.answer, c.tags, r.due_date, r.interval_days, r.ease
        FROM cards c
        JOIN reviews r ON r.card_id = c.id
        WHERE c.unit_id = ? AND c.event_kind = ? AND r.due_date <= ?
        ORDER BY r.due_date ASC
        LIMIT 30;
    """, (unit_id, event_kind, date.today().isoformat())).fetchall()
    con.close()
    return {"cards": [dict(r) for r in rows]}


def sm2_update(interval_days: int, ease: float, grade: int):
    if grade == 0:
        return 1, max(1.3, ease - 0.2)
    if grade == 1:
        return max(1, int(interval_days * 1.2)), max(1.3, ease - 0.15)
    if grade == 2:
        return max(1, int(interval_days * ease)), ease
    return max(1, int(interval_days * (ease + 0.3))), min(3.0, ease + 0.05)


@app.post("/api/review/grade")
def grade_card(card_id: int = Form(...), grade: int = Form(...)):
    from datetime import timedelta
    con = connect()
    row = con.execute("SELECT interval_days, ease FROM reviews WHERE card_id = ?;", (card_id,)).fetchone()
    if not row:
        con.close()
        return JSONResponse({"ok": False, "error": "No review row"}, status_code=404)

    new_interval, new_ease = sm2_update(int(row["interval_days"]), float(row["ease"]), int(grade))
    new_due = (date.today() + timedelta(days=new_interval)).isoformat()

    con.execute("""
      UPDATE reviews
      SET due_date = ?, interval_days = ?, ease = ?, last_grade = ?, updated_at = CURRENT_TIMESTAMP
      WHERE card_id = ?;
    """, (new_due, new_interval, new_ease, int(grade), card_id))
    con.commit()
    con.close()
    return {"ok": True, "next_due": new_due}

# -------------------------
# Quiz endpoints (AE1 practice)
# -------------------------
def _get_tag_weakness(con, table_attempts: str, table_questions: str, unit_id: int, score_col: str = "correct", window: int = 50) -> dict:
    """Compute per-tag weakness scores (0.0=perfect, 1.0=always wrong) from recent attempts."""
    if score_col == "correct":
        query = f"""
            SELECT a.{score_col}, q.tags
            FROM {table_attempts} a
            JOIN {table_questions} q ON q.id = a.question_id
            WHERE a.unit_id = ?
            ORDER BY a.ts DESC LIMIT ?;
        """
    else:
        query = f"""
            SELECT a.score, q.tags
            FROM {table_attempts} a
            JOIN {table_questions} q ON q.id = a.question_id
            WHERE a.unit_id = ?
            ORDER BY a.ts DESC LIMIT ?;
        """
    rows = con.execute(query, (unit_id, window)).fetchall()
    tag_map = {}
    for r in rows:
        tags = [t.strip() for t in (r["tags"] or "").split(",") if t.strip()]
        if not tags:
            continue
        if score_col == "correct":
            val = int(r["correct"])
        else:
            val = 1.0 if int(r["score"]) == 2 else 0.5 if int(r["score"]) == 1 else 0.0
        for t in tags:
            if t not in tag_map:
                tag_map[t] = [0.0, 0]
            tag_map[t][0] += val
            tag_map[t][1] += 1
    # Return weakness: 1.0 - pct_correct; higher = weaker
    weakness = {}
    for tag, (c, n) in tag_map.items():
        weakness[tag] = 1.0 - (c / n) if n > 0 else 0.5
    return weakness

try:
    SYDNEY_TZ = ZoneInfo("Australia/Sydney")
except Exception:
    SYDNEY_TZ = timezone(timedelta(hours=10))


QUIZ_REPEAT_WINDOW_SECONDS = 6 * 60 * 60  # 6 hours without repeats


def _sydney_date_iso() -> str:
    return datetime.now(SYDNEY_TZ).date().isoformat()


def _now_sydney_iso() -> str:
    return datetime.now(SYDNEY_TZ).isoformat()


def _prune_recent_served(con, unit_id: int) -> list:
    cutoff = datetime.now(SYDNEY_TZ) - timedelta(seconds=QUIZ_REPEAT_WINDOW_SECONDS)
    cutoff_iso = cutoff.isoformat()
    con.execute("DELETE FROM quiz_served_history WHERE unit_id = ? AND served_ts < ?;", (unit_id, cutoff_iso))
    rows = con.execute(
        "SELECT question_id FROM quiz_served_history WHERE unit_id = ? AND served_ts >= ?;",
        (unit_id, cutoff_iso),
    ).fetchall()
    return [int(r["question_id"]) for r in rows]


def _record_served_question(con, unit_id: int, question_id: int):
    con.execute(
        "INSERT INTO quiz_served_history (unit_id, question_id, served_ts) VALUES (?, ?, ?);",
        (unit_id, question_id, _now_sydney_iso()),
    )


def _record_daily_progress(
    con,
    unit_id: int,
    quiz_correct: int = 0,
    quiz_total: int = 0,
    sa_correct: int = 0,
    sa_total: int = 0,
) -> None:
    if not (quiz_correct or quiz_total or sa_correct or sa_total):
        return
    dt = _sydney_date_iso()
    con.execute(
        """
        INSERT INTO daily_progress (unit_id, dt, quiz_correct, quiz_total, sa_correct, sa_total)
        VALUES (?, ?, ?, ?, ?, ?)
        ON CONFLICT(unit_id, dt) DO UPDATE SET
          quiz_correct = quiz_correct + excluded.quiz_correct,
          quiz_total = quiz_total + excluded.quiz_total,
          sa_correct = sa_correct + excluded.sa_correct,
          sa_total = sa_total + excluded.sa_total;
        """,
        (unit_id, dt, quiz_correct, quiz_total, sa_correct, sa_total),
    )


def _weighted_pick(candidates: list, tag_weakness: dict) -> dict:
    """Pick a candidate question weighted by tag weakness. Unseen questions get top priority."""
    if not candidates:
        return None
    if len(candidates) == 1:
        return candidates[0]

    weights = []
    for c in candidates:
        tags = [t.strip() for t in (c["tags"] or "").split(",") if t.strip()]
        unseen = c["last_seen"] is None
        if unseen:
            w = 3.0  # unseen questions still prioritized but not dominant
        else:
            # Average weakness of this question's tags; default 0.5 for unknown tags
            tag_scores = [tag_weakness.get(t, 0.5) for t in tags] if tags else [0.5]
            w = max(0.1, sum(tag_scores) / len(tag_scores) * 3.0 + 0.5)
        weights.append(w)

    return random.choices(candidates, weights=weights, k=1)[0]


@app.get("/api/quiz/random")
def quiz_random(unit_id: int, exclude_ids: Optional[str] = Query(None)):
    con = connect()

    total = con.execute(
        "SELECT COUNT(*) AS n FROM quiz_questions WHERE unit_id = ?;", (unit_id,)
    ).fetchone()["n"]

    if total == 0:
        con.close()
        return {"question_id": None, "total": 0}

    # Fetch all questions with last-seen info, pick from a randomised pool
    all_qs = con.execute("""
        SELECT q.id, q.question, q.choices_json, q.context, q.tags,
               MAX(a.ts) AS last_seen
        FROM quiz_questions q
        LEFT JOIN quiz_attempts a ON a.question_id = q.id AND a.unit_id = q.unit_id
        WHERE q.unit_id = ?
        GROUP BY q.id;
    """, (unit_id,)).fetchall()

    all_qs = [dict(r) for r in all_qs]

    exclude_set = set()
    if exclude_ids:
        for part in exclude_ids.split(","):
            part = part.strip()
            if part.isdigit():
                exclude_set.add(int(part))

    filtered = [q for q in all_qs if q["id"] not in exclude_set]
    candidates_pool = filtered if filtered else all_qs
    recent_ids = _prune_recent_served(con, unit_id)
    fresh_candidates = [q for q in candidates_pool if q["id"] not in recent_ids]
    candidates = fresh_candidates if fresh_candidates else candidates_pool
    random.shuffle(candidates)

    # Get per-tag weakness and pick weighted from the full pool
    tag_weakness = _get_tag_weakness(con, "quiz_attempts", "quiz_questions", unit_id, "correct")
    r = _weighted_pick(candidates, tag_weakness)
    if not r:
        con.close()
        return {"question_id": None, "total": total}

    _record_served_question(con, unit_id, r["id"])
    con.close()

    return {
        "question_id": r["id"],
        "question": r["question"],
        "choices": json.loads(r["choices_json"]),
        "context": r["context"] or "",
        "total": total,
    }


def _maybe_create_card(con, unit_id: int, prompt: str, answer: str, tags: str) -> bool:
    """Create a flashcard from a mistake if one doesn't already exist. Returns True if created."""
    existing = con.execute(
        "SELECT id FROM cards WHERE unit_id = ? AND prompt = ?;", (unit_id, prompt.strip())
    ).fetchone()
    if existing:
        return False
    cur = con.cursor()
    cur.execute(
        "INSERT INTO cards (unit_id, event_kind, prompt, answer, tags) VALUES (?, ?, ?, ?, ?);",
        (unit_id, "Knowledge", prompt.strip(), answer.strip(), (tags.strip() + ",auto").strip(","))
    )
    card_id = cur.lastrowid
    cur.execute(
        "INSERT OR REPLACE INTO reviews (card_id, due_date, interval_days, ease) VALUES (?, ?, ?, ?);",
        (card_id, date.today().isoformat(), 1, 2.5)
    )
    return True

@app.post("/api/quiz/answer")
def quiz_answer(unit_id: int = Form(...), question_id: int = Form(...), chosen_index: int = Form(...)):
    con = connect()
    q = con.execute("""
        SELECT question, answer_index, explanation, tags
        FROM quiz_questions
        WHERE id = ? AND unit_id = ?;
    """, (question_id, unit_id)).fetchone()
    if not q:
        con.close()
        return JSONResponse({"ok": False, "error": "Question not found"}, status_code=404)

    correct = 1 if int(chosen_index) == int(q["answer_index"]) else 0
    con.execute("""
        INSERT INTO quiz_attempts (unit_id, question_id, chosen_index, correct, source)
        VALUES (?, ?, ?, ?, ?);
    """, (unit_id, question_id, int(chosen_index), correct, "mcq_quiz"))

    card_created = False
    if correct == 0:
        card_created = _maybe_create_card(con, unit_id, q["question"], q["explanation"], q["tags"] or "")
    _record_daily_progress(con, unit_id, quiz_correct=correct, quiz_total=1)

    con.commit()
    con.close()
    return {"ok": True, "correct": bool(correct), "explanation": q["explanation"], "card_created": card_created}
@app.get("/api/quiz/stats")
def quiz_stats(unit_id: int, window: int = 20, tag_window: int = 50):
    con = connect()

    # Overall stats (last N attempts)
    rows = con.execute("""
        SELECT correct
        FROM quiz_attempts
        WHERE unit_id = ?
        ORDER BY ts DESC
        LIMIT ?;
    """, (unit_id, window)).fetchall()

    total = len(rows)
    correct_n = sum(int(r["correct"]) for r in rows) if rows else 0
    pct = round((correct_n / total) * 100, 1) if total else 0.0

    # Tag breakdown (last tag_window attempts)
    tag_rows = con.execute("""
        SELECT qa.correct, qq.tags
        FROM quiz_attempts qa
        JOIN quiz_questions qq ON qq.id = qa.question_id
        WHERE qa.unit_id = ?
        ORDER BY qa.ts DESC
        LIMIT ?;
    """, (unit_id, tag_window)).fetchall()

    tag_map = {}  # tag -> [correct_sum, total]
    for r in tag_rows:
        tags = (r["tags"] or "").split(",")
        tags = [t.strip() for t in tags if t.strip()]
        if not tags:
            tags = ["(untagged)"]
        for t in tags:
            if t not in tag_map:
                tag_map[t] = [0, 0]
            tag_map[t][0] += int(r["correct"])
            tag_map[t][1] += 1

    by_tag = []
    for tag, (c, n) in tag_map.items():
        by_tag.append({
            "tag": tag,
            "correct": c,
            "total": n,
            "pct": round((c / n) * 100, 1) if n else 0.0
        })

    # Sort weakest first, then by sample size desc
    by_tag.sort(key=lambda x: (x["pct"], -x["total"], x["tag"]))

    con.close()
    return {
        "window": window,
        "total": total,
        "correct": correct_n,
        "pct": pct,
        "by_tag": by_tag[:10]  # top 10 weakest tags
    }

# -------------------------
# Progress dashboard
# -------------------------
@app.get("/api/progress")
def get_progress(unit_id: int, days: int = 14):
    con = connect()
    today = date.today()

    # Get daily quiz stats
    quiz_daily = con.execute("""
        SELECT DATE(ts) AS d, SUM(correct) AS c, COUNT(*) AS n
        FROM quiz_attempts
        WHERE unit_id = ?
        GROUP BY DATE(ts)
        ORDER BY d DESC;
    """, (unit_id,)).fetchall()

    # Get daily short answer stats
    sa_daily = con.execute("""
        SELECT DATE(ts) AS d,
               SUM(CASE WHEN score = 2 THEN 1.0 WHEN score = 1 THEN 0.5 ELSE 0.0 END) AS c,
               COUNT(*) AS n
        FROM short_answer_attempts
        WHERE unit_id = ?
        GROUP BY DATE(ts)
        ORDER BY d DESC;
    """, (unit_id,)).fetchall()

    con.close()

    # Merge into per-day totals
    day_map = {}
    for r in quiz_daily:
        d = r["d"]
        if d not in day_map:
            day_map[d] = [0.0, 0]
        day_map[d][0] += float(r["c"])
        day_map[d][1] += int(r["n"])
    for r in sa_daily:
        d = r["d"]
        if d not in day_map:
            day_map[d] = [0.0, 0]
        day_map[d][0] += float(r["c"])
        day_map[d][1] += int(r["n"])

    # Build last N days array
    result_days = []
    for i in range(days - 1, -1, -1):
        d = (today - timedelta(days=i)).isoformat()
        c, n = day_map.get(d, (0, 0))
        result_days.append({"date": d, "correct": round(c, 1), "total": n})

    # Calculate streak (consecutive days ending today with activity)
    streak = 0
    for i in range(0, 365):
        d = (today - timedelta(days=i)).isoformat()
        if d in day_map and day_map[d][1] > 0:
            streak += 1
        else:
            break

    total_answered = sum(v[1] for v in day_map.values())
    total_correct = sum(v[0] for v in day_map.values())
    overall_pct = round((total_correct / total_answered) * 100, 1) if total_answered else 0.0

    return {
        "days": result_days,
        "streak": streak,
        "total_answered": total_answered,
        "overall_pct": overall_pct,
    }

# -------------------------
# Gemini API key settings
# -------------------------
@app.post("/api/settings/gemini-key")
def save_gemini_key(gemini_api_key: str = Form("")):
    con = connect()
    con.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?);", ("gemini_api_key", gemini_api_key.strip()))
    con.commit()
    con.close()
    return {"ok": True}

@app.get("/api/settings/gemini-key-status")
def gemini_key_status():
    con = connect()
    key = _get_setting(con, "gemini_api_key")
    con.close()
    return {"has_key": bool(key), "key_length": len(key) if key else 0}

@app.get("/api/settings/gemini-key-test")
def gemini_key_test():
    """Quick test that the saved Gemini API key works."""
    con = connect()
    key = _get_setting(con, "gemini_api_key")
    con.close()
    if not key:
        return {"ok": False, "error": "No API key saved"}
    result = _call_gemini(key, "What is 1+1?", "2", "2")
    if result["score"] == -1:
        return {"ok": False, "error": result["feedback"]}
    return {"ok": True, "message": "API key works!", "test_result": result}

# -------------------------
# Short answer quiz endpoints
# -------------------------
@app.get("/api/short-answer/random")
def sa_random(unit_id: int):
    con = connect()

    total = con.execute(
        "SELECT COUNT(*) AS n FROM short_answer_questions WHERE unit_id = ?;", (unit_id,)
    ).fetchone()["n"]

    if total == 0:
        con.close()
        return {"question_id": None, "total": 0}

    all_qs = con.execute("""
        SELECT q.id, q.question, q.context, q.tags,
               MAX(a.ts) AS last_seen
        FROM short_answer_questions q
        LEFT JOIN short_answer_attempts a ON a.question_id = q.id AND a.unit_id = q.unit_id
        WHERE q.unit_id = ?
        GROUP BY q.id;
    """, (unit_id,)).fetchall()

    candidates = [dict(r) for r in all_qs]
    recent_ids = _prune_recent_served(con, unit_id)
    filtered = [q for q in candidates if q["id"] not in recent_ids]
    candidates = filtered if filtered else candidates
    random.shuffle(candidates)
    tag_weakness = _get_tag_weakness(con, "short_answer_attempts", "short_answer_questions", unit_id, "score")

    r = _weighted_pick(candidates, tag_weakness)
    if not r:
        con.close()
        return {"question_id": None, "total": total}

    _record_served_question(con, unit_id, r["id"])
    con.close()

    return {
        "question_id": r["id"],
        "question": r["question"],
        "context": r["context"] or "",
        "total": total,
    }


def _call_gemini(api_key: str, question: str, model_answer: str, student_answer: str) -> dict:
    """Call Gemini API to grade a short answer. Retries on 429 rate-limit."""
    prompt = f"""You are grading a Cert IV networking student's short answer.

Question: {question}
Model answer: {model_answer}
Student's answer: {student_answer}

Rate the answer as: correct, partial, or incorrect.
- "correct" means the student covered the key points adequately.
- "partial" means some key points were covered but important ones are missing.
- "incorrect" means the answer is wrong or misses the point entirely.

Provide brief feedback (1-2 sentences) explaining what was good or missing.

Respond ONLY with valid JSON: {{"score": "correct"|"partial"|"incorrect", "feedback": "..."}}"""

    body = json.dumps({
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.2, "maxOutputTokens": 256}
    }).encode("utf-8")

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"

    # Build SSL context — use system certs, fall back to unverified if needed (Windows compat)
    try:
        ctx = ssl.create_default_context()
    except Exception:
        ctx = ssl._create_unverified_context()

    # Retry up to 3 times on rate-limit (429), waiting between attempts
    max_retries = 3
    for attempt in range(max_retries):
        req = urllib.request.Request(url, data=body, headers={"Content-Type": "application/json"}, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=20, context=ctx) as resp:
                result = json.loads(resp.read().decode("utf-8"))
            text = result["candidates"][0]["content"]["parts"][0]["text"].strip()
            # Strip markdown code fences if present
            if text.startswith("```"):
                text = text.split("\n", 1)[-1]
                if text.endswith("```"):
                    text = text[:-3]
                text = text.strip()
            parsed = json.loads(text)
            score_map = {"correct": 2, "partial": 1, "incorrect": 0}
            return {
                "score": score_map.get(parsed.get("score", "incorrect"), 0),
                "feedback": parsed.get("feedback", "")
            }
        except urllib.error.HTTPError as e:
            if e.code == 429 and attempt < max_retries - 1:
                wait = (attempt + 1) * 4  # 4s, 8s
                print(f"[Gemini] Rate limited (429), retrying in {wait}s (attempt {attempt+1}/{max_retries})...")
                time.sleep(wait)
                continue
            if e.code == 429:
                return {"score": -1, "feedback": "Rate limited after retries — wait a minute and try again."}
            return {"score": -1, "feedback": f"HTTP {e.code}: {e.reason}"}
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {"score": -1, "feedback": str(e)}

    return {"score": -1, "feedback": "Failed after retries."}


def _call_gemini_clarify(
    api_key: str,
    question: str,
    correct_answer: str,
    student_answer: str,
    challenge: str,
    explanation: str,
) -> dict:
    """Ask Gemini to clarify why the correct answer is right based on the student's challenge."""
    student_text = student_answer or "(no student answer provided)"
    prompt = f"""You are a Cert IV networking tutor. Provide a clear explanation that addresses the student's confusion.

Question: {question}
Correct answer: {correct_answer}
Student's answer: {student_text}
Student's clarification request: {challenge}
Teacher notes: {explanation or '(none provided)'}

Explain why the correct answer is right, mention the key concepts that make it correct, and compare it briefly to the student's answer if relevant.

Respond ONLY with valid JSON: {{"clarification": "..."}}"""

    body = json.dumps({
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.2, "maxOutputTokens": 512}
    }).encode("utf-8")

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"

    try:
        ctx = ssl.create_default_context()
    except Exception:
        ctx = ssl._create_unverified_context()

    max_retries = 3
    for attempt in range(max_retries):
        req = urllib.request.Request(url, data=body, headers={"Content-Type": "application/json"}, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=20, context=ctx) as resp:
                result = json.loads(resp.read().decode("utf-8"))
            text = result["candidates"][0]["content"]["parts"][0]["text"].strip()
            if text.startswith("```"):
                text = text.split("\n", 1)[-1]
                if text.endswith("```"):
                    text = text[:-3]
                text = text.strip()
            parsed = None
            try:
                parsed = json.loads(text)
            except json.JSONDecodeError:
                try:
                    sanitized = text.replace("\r", "\\r").replace("\n", "\\n")
                    parsed = json.loads(sanitized)
                except json.JSONDecodeError:
                    m = re.search(r'"clarification"\s*:\s*"(.*)', text, re.DOTALL)
                    if m:
                        raw = m.group(1).replace("\\n", "\n").replace("\\r", "\r").replace('\\"', '"')
                        raw = raw.rstrip('"}').strip()
                        if raw:
                            parsed = {"clarification": raw}
            if parsed and parsed.get("clarification"):
                return {"clarification": parsed.get("clarification", "").strip()}
            return {"clarification": "", "error": "Could not parse clarification from response."}
        except urllib.error.HTTPError as e:
            if e.code == 429 and attempt < max_retries - 1:
                wait = (attempt + 1) * 4
                print(f"[Gemini] Clarify rate limited (429), retrying in {wait}s (attempt {attempt+1}/{max_retries})...")
                time.sleep(wait)
                continue
            if e.code == 429:
                return {"clarification": "", "error": "Rate limited after retries — try again shortly."}
            return {"clarification": "", "error": f"HTTP {e.code}: {e.reason}"}
        except Exception as e:
            return {"clarification": "", "error": str(e)}

    return {"clarification": "", "error": "Failed after retries."}


def _call_gemini_teach(api_key: str, question: str, model_answer: str) -> str:
    """Call Gemini to generate a teaching explanation. Returns plain text."""
    prompt = f"""A Cert IV networking student doesn't know the answer to this question.
Teach them the concept in 3-4 clear sentences, suitable for a beginner.
Use simple language and a concrete example if possible.

Question: {question}
Key points to cover: {model_answer}

Respond in plain text (no JSON, no markdown formatting)."""

    body = json.dumps({
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.4, "maxOutputTokens": 512}
    }).encode("utf-8")

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"

    try:
        ctx = ssl.create_default_context()
    except Exception:
        ctx = ssl._create_unverified_context()

    max_retries = 3
    for attempt in range(max_retries):
        req = urllib.request.Request(url, data=body, headers={"Content-Type": "application/json"}, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=20, context=ctx) as resp:
                result = json.loads(resp.read().decode("utf-8"))
            return result["candidates"][0]["content"]["parts"][0]["text"].strip()
        except urllib.error.HTTPError as e:
            if e.code == 429 and attempt < max_retries - 1:
                time.sleep((attempt + 1) * 4)
                continue
            return ""
        except Exception:
            return ""
    return ""

@app.post("/api/short-answer/teach")
def sa_teach(unit_id: int = Form(...), question_id: int = Form(...)):
    con = connect()
    q = con.execute("""
        SELECT question, model_answer, explanation, tags
        FROM short_answer_questions
        WHERE id = ? AND unit_id = ?;
    """, (question_id, unit_id)).fetchone()
    if not q:
        con.close()
        return JSONResponse({"ok": False, "error": "Question not found"}, status_code=404)

    api_key = _get_setting(con, "gemini_api_key")

    teaching = ""
    if api_key:
        teaching = _call_gemini_teach(api_key, q["question"], q["model_answer"])

    # If no API key or Gemini failed, use the stored explanation as fallback
    if not teaching:
        teaching = q["explanation"]

    # Record attempt as score=0
    con.execute("""
        INSERT INTO short_answer_attempts (unit_id, question_id, student_answer, score, ai_feedback, source)
        VALUES (?, ?, ?, ?, ?, ?);
    """, (unit_id, question_id, "", 0, "Used Teach Me", "short_answer"))

    card_created = _maybe_create_card(con, unit_id, q["question"], q["model_answer"], q["tags"] or "")
    _record_daily_progress(con, unit_id, sa_correct=0, sa_total=1)

    con.commit()
    con.close()

    return {
        "ok": True,
        "teaching": teaching,
        "model_answer": q["model_answer"],
        "card_created": card_created,
    }

@app.post("/api/short-answer/check")
def sa_check(unit_id: int = Form(...), question_id: int = Form(...), student_answer: str = Form(...)):
    con = connect()
    q = con.execute("""
        SELECT question, model_answer, explanation
        FROM short_answer_questions
        WHERE id = ? AND unit_id = ?;
    """, (question_id, unit_id)).fetchone()
    if not q:
        con.close()
        return JSONResponse({"ok": False, "error": "Question not found"}, status_code=404)

    api_key = _get_setting(con, "gemini_api_key")

    if not api_key:
        # Self-grade mode: return model answer, don't store attempt yet
        con.close()
        return {
            "ok": True,
            "mode": "self_grade",
            "model_answer": q["model_answer"],
        }

    # AI grading
    result = _call_gemini(api_key, q["question"], q["model_answer"], student_answer.strip())

    if result["score"] == -1:
        # AI call failed — fall back to self-grade
        con.close()
        return {
            "ok": True,
            "mode": "self_grade",
            "model_answer": q["model_answer"],
            "ai_error": result["feedback"]
        }

    con.execute("""
        INSERT INTO short_answer_attempts (unit_id, question_id, student_answer, score, ai_feedback, source)
        VALUES (?, ?, ?, ?, ?, ?);
    """, (unit_id, question_id, student_answer.strip(), result["score"], result["feedback"], "short_answer"))

    card_created = False
    if result["score"] == 0:
        tags_row = con.execute("SELECT tags FROM short_answer_questions WHERE id = ?;", (question_id,)).fetchone()
        card_created = _maybe_create_card(con, unit_id, q["question"], q["model_answer"], (tags_row["tags"] if tags_row else "") or "")
    _record_daily_progress(con, unit_id, sa_correct=1 if result["score"] == 2 else 0, sa_total=1)

    con.commit()
    con.close()

    return {
        "ok": True,
        "mode": "ai_graded",
        "score": result["score"],
        "ai_feedback": result["feedback"],
        "model_answer": q["model_answer"],
        "card_created": card_created,
    }

@app.post("/api/short-answer/self-grade")
def sa_self_grade(unit_id: int = Form(...), question_id: int = Form(...), score: int = Form(...), student_answer: str = Form("")):
    con = connect()
    actual_score = max(0, min(2, score))
    con.execute("""
        INSERT INTO short_answer_attempts (unit_id, question_id, student_answer, score, ai_feedback, source)
        VALUES (?, ?, ?, ?, ?, ?);
    """, (unit_id, question_id, student_answer.strip(), actual_score, "Self-graded", "short_answer"))

    card_created = False
    if actual_score == 0:
        q = con.execute("SELECT question, model_answer, tags FROM short_answer_questions WHERE id = ?;", (question_id,)).fetchone()
        if q:
            card_created = _maybe_create_card(con, unit_id, q["question"], q["model_answer"], q["tags"] or "")
    _record_daily_progress(con, unit_id, sa_correct=1 if actual_score == 2 else 0, sa_total=1)

    con.commit()
    con.close()
    return {"ok": True, "card_created": card_created}

@app.get("/api/short-answer/stats")
def sa_stats(unit_id: int, window: int = 20, tag_window: int = 50):
    con = connect()

    rows = con.execute("""
        SELECT score
        FROM short_answer_attempts
        WHERE unit_id = ?
        ORDER BY ts DESC
        LIMIT ?;
    """, (unit_id, window)).fetchall()

    total = len(rows)
    # Score 2 = correct, score 1 = partial (count as 0.5), score 0 = incorrect
    correct_n = sum(1 for r in rows if int(r["score"]) == 2)
    partial_n = sum(1 for r in rows if int(r["score"]) == 1)
    pct = round(((correct_n + partial_n * 0.5) / total) * 100, 1) if total else 0.0

    tag_rows = con.execute("""
        SELECT sa.score, sq.tags
        FROM short_answer_attempts sa
        JOIN short_answer_questions sq ON sq.id = sa.question_id
        WHERE sa.unit_id = ?
        ORDER BY sa.ts DESC
        LIMIT ?;
    """, (unit_id, tag_window)).fetchall()

    tag_map = {}
    for r in tag_rows:
        tags = (r["tags"] or "").split(",")
        tags = [t.strip() for t in tags if t.strip()]
        if not tags:
            tags = ["(untagged)"]
        for t in tags:
            if t not in tag_map:
                tag_map[t] = [0, 0]
            tag_map[t][0] += 1 if int(r["score"]) == 2 else 0.5 if int(r["score"]) == 1 else 0
            tag_map[t][1] += 1

    by_tag = []
    for tag, (c, n) in tag_map.items():
        by_tag.append({
            "tag": tag,
            "correct": c,
            "total": n,
            "pct": round((c / n) * 100, 1) if n else 0.0
        })

    by_tag.sort(key=lambda x: (x["pct"], -x["total"], x["tag"]))
    con.close()

    return {
        "window": window,
        "total": total,
        "correct": correct_n,
        "pct": pct,
        "by_tag": by_tag[:10]
    }


# -------------------------
# Mistake review endpoint
# -------------------------
@app.get("/api/mistakes")
def mistakes(unit_id: int):
    """Return questions the student has answered incorrectly (MCQ + short answer)."""
    con = connect()

    quiz_rows = con.execute("""
        SELECT
            a.unit_id,
            a.question_id,
            q.question,
            q.choices_json,
            q.answer_index,
            q.explanation,
            q.tags,
            q.context,
            a.chosen_index,
            a.source,
            a.ts,
            (
                SELECT MAX(a2.correct)
                FROM quiz_attempts a2
                WHERE a2.unit_id = a.unit_id AND a2.question_id = a.question_id
            ) AS ever_correct
        FROM quiz_attempts a
        JOIN quiz_questions q ON q.id = a.question_id AND q.unit_id = a.unit_id
        WHERE a.unit_id = ? AND a.correct = 0
        ORDER BY a.ts DESC;
    """, (unit_id,)).fetchall()

    quiz_map = {}
    for r in quiz_rows:
        qid = r["question_id"]
        if qid not in quiz_map:
            quiz_map[qid] = {
                "question_id": qid,
                "question": r["question"],
                "choices": json.loads(r["choices_json"]),
                "correct_index": int(r["answer_index"]),
                "last_chosen_index": int(r["chosen_index"]),
                "tags": r["tags"] or "",
                "context": r["context"] or "",
                "source": r["source"],
                "last_ts": r["ts"],
                "wrong_count": 0,
                "mastered": bool(r["ever_correct"]),
            }
        quiz_map[qid]["wrong_count"] += 1

    sa_rows = con.execute("""
        SELECT
            a.unit_id,
            a.question_id,
            q.question,
            q.model_answer,
            q.tags,
            q.context,
            a.student_answer,
            a.source,
            a.ts,
            (
                SELECT MAX(a2.score)
                FROM short_answer_attempts a2
                WHERE a2.unit_id = a.unit_id AND a2.question_id = a.question_id
            ) AS max_score
        FROM short_answer_attempts a
        JOIN short_answer_questions q ON q.id = a.question_id AND q.unit_id = a.unit_id
        WHERE a.unit_id = ? AND a.score = 0
        ORDER BY a.ts DESC;
    """, (unit_id,)).fetchall()

    sa_map = {}
    for r in sa_rows:
        qid = r["question_id"]
        if qid not in sa_map:
            sa_map[qid] = {
                "question_id": qid,
                "question": r["question"],
                "model_answer": r["model_answer"],
                "context": r["context"] or "",
                "last_student_answer": r["student_answer"],
                "tags": r["tags"] or "",
                "source": r["source"],
                "last_ts": r["ts"],
                "wrong_count": 0,
                "mastered": bool(r["max_score"] and int(r["max_score"]) > 0),
            }
        sa_map[qid]["wrong_count"] += 1

    con.close()
    return {
        "quiz": list(quiz_map.values()),
        "short_answer": list(sa_map.values()),
    }


# -------------------------
# Exam mode (simple in-memory sessions)
# -------------------------
_EXAMS = {}


@app.post("/api/exam/start")
def exam_start(
    unit_id: int = Form(...),
    count: int = Form(20),
    minutes: int = Form(30),
):
    """Start a timed exam session."""
    con = connect()

    # Build candidate pool: MCQ and SA, favour weak tags by reusing existing helpers
    quiz_rows = con.execute(
        "SELECT id, question, choices_json, context, tags FROM quiz_questions WHERE unit_id = ?;",
        (unit_id,),
    ).fetchall()
    sa_rows = con.execute(
        "SELECT id, question, context, tags FROM short_answer_questions WHERE unit_id = ?;",
        (unit_id,),
    ).fetchall()

    if not quiz_rows and not sa_rows:
        con.close()
        return JSONResponse({"ok": False, "error": "No questions available for this unit"}, status_code=400)

    # Basic pool: mix MCQ and SA; front-end just needs ids and kind
    pool = []
    for r in quiz_rows:
        pool.append(
            {
                "id": r["id"],
                "kind": "mcq",
                "question": r["question"],
                "choices": json.loads(r["choices_json"]),
                "context": r["context"] or "",
                "tags": r["tags"] or "",
            }
        )
    for r in sa_rows:
        pool.append(
            {
                "id": r["id"],
                "kind": "sa",
                "question": r["question"],
                "choices": [],
                "context": r["context"] or "",
                "tags": r["tags"] or "",
            }
        )

    if not pool:
        con.close()
        return JSONResponse({"ok": False, "error": "No questions available for this unit"}, status_code=400)

    random.shuffle(pool)
    count = max(1, min(count, len(pool)))
    questions = pool[:count]

    exam_id = str(int(time.time() * 1000)) + "-" + str(random.randint(1000, 9999))
    now = time.time()
    duration = max(5 * 60, min(minutes * 60, 90 * 60))  # clamp between 5 and 90 min

    _EXAMS[exam_id] = {
        "unit_id": unit_id,
        "created": now,
        "duration": duration,
        "questions": questions,
        "answers": [None] * len(questions),
    }

    con.close()
    return {
        "exam_id": exam_id,
        "questions": questions,
        "current_index": 0,
        "remaining_seconds": duration,
    }


def _exam_get(exam_id: str):
    exam = _EXAMS.get(exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found or expired")
    now = time.time()
    if now - exam["created"] > exam["duration"] + 60:
        _EXAMS.pop(exam_id, None)
        raise HTTPException(status_code=410, detail="Exam has expired")
    return exam, now


def _compute_exam_summary_and_persist(exam_id: str):
    """Compute exam summary, persist attempts to DB (so Mistakes tab shows them), pop exam, return summary."""
    exam = _EXAMS.get(exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found or expired")
    now = time.time()
    con = connect()
    unit_id = exam["unit_id"]
    questions = exam["questions"]
    answers = exam["answers"]
    api_key = _get_setting(con, "gemini_api_key")

    correct = 0
    total = len(questions)
    review = []

    for idx, q in enumerate(questions):
        your = answers[idx] if idx < len(answers) else None
        if q["kind"] == "mcq":
            row = con.execute(
                "SELECT choices_json, answer_index, question, explanation, tags FROM quiz_questions WHERE id = ? AND unit_id = ?;",
                (q["id"], unit_id),
            ).fetchone()
            if row:
                choices = json.loads(row["choices_json"])
                try:
                    chosen_idx = int(your) if your is not None and your != "" else None
                except (ValueError, TypeError):
                    chosen_idx = None
                correct_idx = int(row["answer_index"])
                is_correct = chosen_idx is not None and chosen_idx == correct_idx
                if is_correct:
                    correct += 1
                review.append({
                    "kind": "mcq",
                    "question_id": q["id"],
                    "question": q["question"],
                    "your_answer": choices[chosen_idx] if chosen_idx is not None and 0 <= chosen_idx < len(choices) else "",
                    "correct_answer": choices[correct_idx] if 0 <= correct_idx < len(choices) else "",
                    "correct": is_correct,
                })
                con.execute(
                    "INSERT INTO quiz_attempts (unit_id, question_id, chosen_index, correct, source) VALUES (?, ?, ?, ?, ?);",
                    (unit_id, q["id"], chosen_idx if chosen_idx is not None else -1, 1 if is_correct else 0, "exam"),
                )
                if not is_correct:
                    _maybe_create_card(con, unit_id, row["question"], row["explanation"], row["tags"] or "")
        else:
            row = con.execute(
                "SELECT model_answer, question, explanation, tags FROM short_answer_questions WHERE id = ? AND unit_id = ?;",
                (q["id"], unit_id),
            ).fetchone()
            if row:
                student_ans = (your or "").strip()
                score = 0
                feedback = ""
                if student_ans and api_key:
                    result = _call_gemini(api_key, row["question"], row["model_answer"], student_ans)
                    score = max(0, min(2, result["score"])) if result["score"] >= 0 else 0
                    feedback = result.get("feedback", "")
                if score == 2:
                    correct += 1
                is_correct = score == 2
                review.append({
                    "kind": "sa",
                    "question_id": q["id"],
                    "question": q["question"],
                    "your_answer": your or "",
                    "correct_answer": row["model_answer"],
                    "correct": is_correct,
                    "ai_feedback": feedback,
                })
                con.execute(
                    "INSERT INTO short_answer_attempts (unit_id, question_id, student_answer, score, ai_feedback, source) VALUES (?, ?, ?, ?, ?, ?);",
                    (unit_id, q["id"], student_ans, score, feedback, "exam"),
                )
                if score == 0:
                    _maybe_create_card(con, unit_id, row["question"], row["model_answer"], row["tags"] or "")

    pct = (correct / total) * 100 if total else 0.0
    summary = {
        "correct": correct,
        "total": total,
        "pct": pct,
        "time_used_seconds": int(now - exam["created"]),
        "questions": review,
        "unit_id": unit_id,
    }
    con.commit()
    con.close()
    _EXAMS.pop(exam_id, None)
    return summary


@app.post("/api/exam/answer")
def exam_answer(exam_id: str = Form(...), index: int = Form(...), answer: str = Form("")):
    exam, now = _exam_get(exam_id)
    idx = max(0, min(int(index), len(exam["questions"]) - 1))
    exam["answers"][idx] = answer

    remaining = max(0.0, exam["duration"] - (now - exam["created"]))
    finished = remaining <= 0 or idx >= len(exam["questions"]) - 1

    if finished:
        summary = _compute_exam_summary_and_persist(exam_id)
        return {"exam_id": exam_id, "finished": True, "summary": summary}

    current_index = idx + 1
    return {
        "exam_id": exam_id,
        "questions": exam["questions"],
        "current_index": current_index,
        "remaining_seconds": remaining,
        "finished": False,
    }


@app.post("/api/exam/finish")
def exam_finish(exam_id: str = Form(...)):
    summary = _compute_exam_summary_and_persist(exam_id)
    return {"exam_id": exam_id, "finished": True, "summary": summary}


@app.post("/api/exam/challenge")
def exam_challenge(
    unit_id: int = Form(...),
    question_id: int = Form(...),
    kind: str = Form(...),
    user_answer: str = Form(""),
    user_challenge: str = Form(""),
):
    if not user_challenge or not user_challenge.strip():
        return JSONResponse({"ok": False, "error": "Describe what you want clarified."}, status_code=400)
    con = connect()
    api_key = _get_setting(con, "gemini_api_key")
    if not api_key:
        con.close()
        return JSONResponse({"ok": False, "error": "Gemini API key not configured."}, status_code=400)

    kind_key = (kind or "").lower()
    if kind_key.startswith("mcq"):
        row = con.execute(
            "SELECT question, choices_json, answer_index, explanation FROM quiz_questions WHERE id = ? AND unit_id = ?;",
            (question_id, unit_id),
        ).fetchone()
        if not row:
            con.close()
            return JSONResponse({"ok": False, "error": "Question not found."}, status_code=404)
        choices = json.loads(row["choices_json"])
        try:
            answer_idx = int(row["answer_index"])
        except (ValueError, TypeError):
            answer_idx = None
        correct_answer = choices[answer_idx] if answer_idx is not None and 0 <= answer_idx < len(choices) else row["explanation"]
        try:
            candidate_idx = int(user_answer)
            user_answer_text = choices[candidate_idx] if 0 <= candidate_idx < len(choices) else user_answer
        except ValueError:
            user_answer_text = user_answer or "(no answer provided)"
        explanation = row["explanation"] or ""
        question_text = row["question"]
    else:
        row = con.execute(
            "SELECT question, model_answer, explanation FROM short_answer_questions WHERE id = ? AND unit_id = ?;",
            (question_id, unit_id),
        ).fetchone()
        if not row:
            con.close()
            return JSONResponse({"ok": False, "error": "Question not found."}, status_code=404)
        correct_answer = row["model_answer"]
        user_answer_text = user_answer or "(no answer provided)"
        explanation = row["explanation"] or ""
        question_text = row["question"]

    result = _call_gemini_clarify(api_key, question_text, correct_answer, user_answer_text, user_challenge.strip(), explanation)
    con.close()
    clarification = result.get("clarification", "").strip()
    if clarification:
        return {"ok": True, "clarification": clarification}
    return {"ok": False, "error": result.get("error", "Could not clarify that question right now.")}


# -------------------------
# Explain-it-back endpoints
# -------------------------
@app.get("/api/explain/random")
def explain_random(unit_id: int):
    con = connect()
    rows = con.execute(
        "SELECT id, topic_prompt, model_explanation, tags FROM explain_topics WHERE unit_id = ?;",
        (unit_id,),
    ).fetchall()
    con.close()
    if not rows:
        return {"topic_id": None}
    row = dict(random.choice(rows))
    return {
        "topic_id": row["id"],
        "topic_prompt": row["topic_prompt"],
        "model_explanation": row["model_explanation"],
        "tags": row["tags"] or "",
    }


@app.post("/api/explain/check")
def explain_check(
    unit_id: int = Form(...),
    topic_id: int = Form(...),
    student_text: str = Form(...),
):
    con = connect()
    row = con.execute(
        "SELECT topic_prompt, model_explanation, tags FROM explain_topics WHERE id = ? AND unit_id = ?;",
        (topic_id, unit_id),
    ).fetchone()
    if not row:
        con.close()
        return JSONResponse({"ok": False, "error": "Topic not found"}, status_code=404)

    prompt = row["topic_prompt"]
    model = row["model_explanation"]

    api_key = _get_setting(con, "gemini_api_key")
    score = 3
    feedback = ""

    if api_key:
        rubric = f"""You are marking a Cert IV networking student's explanation.

Topic prompt: {prompt}
Model explanation: {model}
Student explanation: {student_text}

Rate the student explanation from 1 to 5 for:
- accuracy (technical correctness)
- completeness (covers key points from the model explanation)
- clarity (understandable, uses appropriate terminology)

Respond ONLY with valid JSON: {{"score": 1-5, "feedback": "brief comments on what was good and what was missing"}}"""

        body = json.dumps({
            "contents": [{"parts": [{"text": rubric}]}],
            "generationConfig": {"temperature": 0.2, "maxOutputTokens": 256}
        }).encode("utf-8")

        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
        try:
            try:
                ctx = ssl.create_default_context()
            except Exception:
                ctx = ssl._create_unverified_context()
            req = urllib.request.Request(url, data=body, headers={"Content-Type": "application/json"}, method="POST")
            with urllib.request.urlopen(req, timeout=20, context=ctx) as resp:
                result = json.loads(resp.read().decode("utf-8"))
            text = result["candidates"][0]["content"]["parts"][0]["text"].strip()
            if text.startswith("```"):
                text = text.split("\n", 1)[-1]
                if text.endswith("```"):
                    text = text[:-3]
                text = text.strip()
            parsed = json.loads(text)
            s_val = int(parsed.get("score", 3))
            score = max(1, min(5, s_val))
            feedback = parsed.get("feedback", "")
        except Exception:
            # Fall back to default score and generic feedback
            score = 3
            feedback = "AI feedback unavailable. Compare your explanation with the model answer."
    else:
        feedback = "No AI key configured. Compare your explanation with the model answer and check that you cover the key points."

    con.execute(
        """
        INSERT INTO explain_attempts (unit_id, topic_id, student_text, ai_score, ai_feedback)
        VALUES (?, ?, ?, ?, ?);
        """,
        (unit_id, topic_id, student_text.strip(), score, feedback),
    )
    con.commit()
    con.close()

    return {
        "ok": True,
        "ai_score": score,
        "ai_feedback": feedback,
        "model_explanation": model,
    }


# -------------------------
# AE2 endpoints (Case Study builder)
# -------------------------
@app.get("/api/ae2/items")
def ae2_items(unit_id: int):
    con = connect()
    rows = con.execute("""
        SELECT code, title, section, order_index
        FROM ae2_items
        WHERE unit_id = ?
        ORDER BY order_index ASC;
    """, (unit_id,)).fetchall()
    con.close()
    return {"items": [dict(r) for r in rows]}


@app.get("/api/ae2/item")
def ae2_item(unit_id: int, item_code: str):
    con = connect()
    item = con.execute("""
        SELECT code, title, section, template_md, word_guidance
        FROM ae2_items
        WHERE unit_id = ? AND code = ?;
    """, (unit_id, item_code)).fetchone()

    resp = con.execute("""
        SELECT content_md
        FROM ae2_responses
        WHERE unit_id = ? AND item_code = ?;
    """, (unit_id, item_code)).fetchone()

    con.close()
    if not item:
        return JSONResponse({"ok": False, "error": "Item not found"}, status_code=404)

    return {
        "code": item["code"],
        "title": item["title"],
        "section": item["section"],
        "template_md": item["template_md"],
        "word_guidance": item["word_guidance"],
        "content_md": resp["content_md"] if resp else ""
    }


@app.post("/api/ae2/response/save")
def ae2_save(unit_id: int = Form(...), item_code: str = Form(...), content_md: str = Form(...)):
    con = connect()
    con.execute("""
        INSERT OR REPLACE INTO ae2_responses (unit_id, item_code, content_md, updated_at)
        VALUES (?, ?, ?, CURRENT_TIMESTAMP);
    """, (unit_id, item_code.strip(), content_md))
    con.commit()
    con.close()
    return {"ok": True}


def _get_setting(con, key: str) -> str:
    r = con.execute("SELECT value FROM settings WHERE key = ?;", (key,)).fetchone()
    return r["value"] if r else ""


def _md_to_docx(document: Document, md: str):
    lines = md.replace("\r\n", "\n").split("\n")
    i = 0

    def is_table_line(s: str) -> bool:
        s = s.strip()
        return s.startswith("|") and s.endswith("|") and s.count("|") >= 2

    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("#"):
            level = min(4, len(line) - len(line.lstrip("#")))
            text = line.lstrip("#").strip()
            document.add_heading(text, level=level)
            i += 1
            continue

        # Bullet
        if line.strip().startswith("- "):
            document.add_paragraph(line.strip()[2:].strip(), style="List Bullet")
            i += 1
            continue

        # Table (pipe markdown)
        if is_table_line(line) and i + 1 < len(lines) and ("---" in lines[i + 1]):
            table_lines = []
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i].strip())
                i += 1

            rows = [[c.strip() for c in tl.strip("|").split("|")] for tl in table_lines]
            if len(rows) >= 2:
                header = rows[0]
                body = [r for r in rows[2:]]  # skip separator row

                cols = max(len(header), *(len(r) for r in body)) if body else len(header)
                t = document.add_table(rows=1 + len(body), cols=cols)
                t.style = "Table Grid"

                for c in range(cols):
                    t.cell(0, c).text = header[c] if c < len(header) else ""

                for r_idx, r in enumerate(body, start=1):
                    for c in range(cols):
                        t.cell(r_idx, c).text = r[c] if c < len(r) else ""

            continue

        # Paragraph (default)
    document.add_paragraph(line)
    i += 1


@app.get("/api/ae2/export_docx")
def ae2_export_docx(unit_id: int, item_code: str):
    con = connect()
    item = con.execute("""
        SELECT title, template_md
        FROM ae2_items
        WHERE unit_id = ? AND code = ?;
    """, (unit_id, item_code)).fetchone()

    resp = con.execute("""
        SELECT content_md
        FROM ae2_responses
        WHERE unit_id = ? AND item_code = ?;
    """, (unit_id, item_code)).fetchone()

    if not item:
        con.close()
        return JSONResponse({"ok": False, "error": "Item not found"}, status_code=404)

    student_name = _get_setting(con, "student_name") or "Student"
    student_number = _get_setting(con, "student_number") or ""
    con.close()

    content = (resp["content_md"] if resp else "").strip()
    if not content:
        content = item["template_md"]

    doc = Document()
    doc.add_heading("ICTNWK421/423 – Assessment Event 2 (Case Study)", level=1)
    doc.add_paragraph(f"Item: {item_code} — {item['title']}")
    doc.add_paragraph("")

    _md_to_docx(doc, content)

    # Footer text uses the first footer paragraph (python-docx model) [web:336][web:338]
    footer_text = student_name + (f" | {student_number}" if student_number else "")
    section = doc.sections[0]
    footer = section.footer
    if footer.paragraphs:
        footer.paragraphs[0].text = footer_text
    else:
        footer.add_paragraph(footer_text)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    safe_code = item_code.replace("/", "_").replace("\\", "_").replace(" ", "_")
    filename = f"AE2_{safe_code}.docx"

    # Streaming download with proper content-disposition [web:341][web:350]
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        bio,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )
